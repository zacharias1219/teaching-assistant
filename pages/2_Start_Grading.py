# 2_Start_Grading.py (FIXED: Text wrapping for remarks)

import streamlit as st
import fitz
import json
import openai
from PIL import Image
import io
import re
import math
import pytesseract
import os
from datetime import datetime
import toml
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Start Grading", layout="wide")

# Initialize session state for prompts
if "grading_prompt" not in st.session_state:
    st.session_state.grading_prompt = """You are an expert mathematics professor grading multiple questions. Analyze the student's work for ALL questions below.

**QUESTIONS TO GRADE:**
{questions_text}

**GRADING CRITERIA:**
- **IGNORE crossed-out content** - only grade final, uncrossed work
- **Check completion** - penalize incomplete solutions (e.g., "100x-50x" without "=50x")
- **EXCELLENT (90-100%)**: Complete, correct, proper notation, clear reasoning
- **GOOD (70-89%)**: Mostly correct, minor errors, nearly complete
- **FAIR (50-69%)**: Partial understanding, significant errors, incomplete
- **POOR (20-49%)**: Major errors, wrong approach, abandoned work
- **NOT ATTEMPTED (0%)**: No relevant work

**TASK:** Grade each question by finding the relevant work in the images and evaluating according to the criteria above.

**REQUIRED JSON OUTPUT:**
{{
    "question_analysis": [
        {{
            "question_number": <question_number>,
            "status": "<'Excellent'|'Good'|'Fair'|'Poor'|'Not Attempted'>",
            "score": <integer between 0 and max_score>,
            "max_score": <max_score>,
            "feedback": "<brief feedback covering: strengths, errors, completion status, and improvement suggestions>",
            "extracted_work": "<description of work found for this question>",
            "mathematical_quality": "<assessment of rigor and notation>",
            "completion_status": "<'Complete'|'Nearly Complete'|'Partially Complete'|'Incomplete'|'Abandoned'>"
        }}
    ]
}}"""

if "summary_prompt" not in st.session_state:
    st.session_state.summary_prompt = """You are a senior mathematics education specialist with expertise in curriculum development and student assessment. You are creating a comprehensive academic evaluation report for a student's mathematical performance.

**STUDENT'S DETAILED TEST RESULTS:**
{full_analysis}

**CURRICULUM STANDARDS & RUBRIC:**
"{rubric}"

**ACADEMIC EVALUATION REQUIREMENTS:**

**1. OVERALL PERFORMANCE ASSESSMENT:**
Create a comprehensive academic summary (3-4 sentences) that includes:
- **Performance Level**: Overall grade/performance classification
- **Academic Strengths**: Specific mathematical competencies demonstrated
- **Learning Gaps**: Identified areas of conceptual weakness
- **Growth Indicators**: Evidence of mathematical thinking development
- **Strategic Recommendations**: Specific, actionable study and practice suggestions

**2. CONCEPTUAL MASTERY ANALYSIS:**
For each core mathematical concept, provide:
- **Concept Identification**: Clear definition of the mathematical concept
- **Performance Assessment**: Detailed evaluation of understanding level
- **Evidence-Based Analysis**: Specific examples from student work
- **Learning Progression**: Where the student stands in concept mastery
- **Next Steps**: Specific recommendations for concept development

**3. MATHEMATICAL THINKING EVALUATION:**
Assess the student's mathematical reasoning abilities:
- **Problem-Solving Approach**: How they tackle mathematical challenges
- **Logical Reasoning**: Quality of mathematical argumentation
- **Notational Proficiency**: Command of mathematical language
- **Computational Accuracy**: Precision in calculations
- **Conceptual Connections**: Ability to link related mathematical ideas

**4. EDUCATIONAL RECOMMENDATIONS:**
Provide specific, actionable guidance for:
- **Immediate Focus Areas**: What to study next
- **Practice Strategies**: How to improve specific skills
- **Concept Reinforcement**: Ways to strengthen understanding
- **Advanced Preparation**: Preparation for next level mathematics

**REQUIRED JSON OUTPUT:**
{{
    "overall_remarks": "<comprehensive academic summary covering performance level, strengths, gaps, and strategic recommendations>",
    "rubric_analysis": [
        {{
            "concept": "<specific mathematical concept from rubric>",
            "performance": "<'Mastery'|'Proficient'|'Developing'|'Emerging'|'Not Demonstrated'>",
            "evidence": "<detailed evidence from student work with specific examples and mathematical reasoning>",
            "recommendations": "<specific, actionable recommendations for this concept>"
        }}
    ],
    "mathematical_thinking": "<assessment of problem-solving approach, logical reasoning, and mathematical communication>",
    "learning_recommendations": "<comprehensive study and practice recommendations for continued mathematical development>"
}}"""

# Add debug mode toggle
debug_mode = st.sidebar.checkbox("🐛 Debug Mode", value=False, help="Enable debug information to troubleshoot issues")
st.session_state.debug_mode = debug_mode

# Force reset the OCR prompt to a high-fidelity full-transcription version
st.session_state.ocr_prompt = """TRANSCRIBE ALL VISIBLE TEXT EXACTLY (FOCUS: COMPLETE HANDWRITTEN ANSWERS)

You are a high-accuracy OCR agent. Capture the ENTIRE written work as plain text, preserving structure.

STRICT RULES
- Preserve ALL content and line breaks; do not omit steps before or after equations; never summarize.
- Keep numbering/bullets exactly as written (e.g., "(i)", "(a)", "1.", "Q4"). Do not renumber or merge lines.
- Preserve math faithfully with readable inline notation: sqrt(), abs(), d/dx, ∫, Σ, Π, lim, |x|, ^ for powers, / for fractions, parentheses for grouping.
- Convert LaTeX-like tokens to readable math (e.g., \\frac{dy}{dx} -> d/dx, \\sqrt{x} -> sqrt(x)). Remove LaTeX backslashes and $.
- Keep symbols like →, ≈, ≤, ≥ when present; if unclear, use ->.
- If any token is unreadable, write [illegible]; do not invent content.
- If diagrams/figures exist, write [diagram] on its own line; for tables, transcribe cell text line by line.
- Keep the original language. No commentary, no extra instructions.

OUTPUT FORMAT
- Plain text only with original line breaks. No JSON. No markdown. No headings."""

# Sidebar for prompt editing
with st.sidebar:
    st.header("⚙️ System Prompts Configuration")
    st.info("Edit these prompts to customize the AI's behavior. Changes take effect immediately.")
    
    # Grading prompt editor
    st.subheader("📝 Grading Prompt")
    st.caption("Controls how the AI grades individual questions")
    new_grading_prompt = st.text_area(
        "Grading System Prompt",
        value=st.session_state.grading_prompt,
        height=400,
        help="This prompt controls the AI's grading behavior for individual questions"
    )
    if new_grading_prompt != st.session_state.grading_prompt:
        st.session_state.grading_prompt = new_grading_prompt
        st.success("✅ Grading prompt updated!")
    
    # Summary prompt editor
    st.subheader("📊 Summary Prompt")
    st.caption("Controls how the AI generates overall summaries")
    new_summary_prompt = st.text_area(
        "Summary System Prompt",
        value=st.session_state.summary_prompt,
        height=300,
        help="This prompt controls the AI's summary generation behavior"
    )
    if new_summary_prompt != st.session_state.summary_prompt:
        st.session_state.summary_prompt = new_summary_prompt
        st.success("✅ Summary prompt updated!")
    
    # OCR prompt editor (now handled per question)
    st.subheader("🔍 OCR Configuration")
    st.caption("OCR prompts are now configured individually for each question")
    st.info("💡 Use the 'Test Questions & OCR Configuration' section below to set custom OCR prompts for each question.")
    
    # Debug mode toggle
    st.divider()
    st.subheader("🔧 Debug Options")
    debug_mode = st.checkbox("Show Debug Information", value=st.session_state.get('debug_mode', False))
    if debug_mode != st.session_state.get('debug_mode', False):
        st.session_state.debug_mode = debug_mode
        st.success("✅ Debug mode updated!")
    
    # Force reset OCR prompt
    if st.button("🔄 Force Reset OCR Prompt"):
        st.session_state.ocr_prompt = """TRANSCRIBE ALL VISIBLE TEXT EXACTLY (FOCUS: COMPLETE HANDWRITTEN ANSWERS)

You are a high-accuracy OCR agent. Capture the ENTIRE written work as plain text, preserving structure.

STRICT RULES
- Preserve ALL content and line breaks; do not omit steps before or after equations; never summarize.
- Keep numbering/bullets exactly as written (e.g., "(i)", "(a)", "1.", "Q4"). Do not renumber or merge lines.
- Preserve math faithfully with readable inline notation: sqrt(), abs(), d/dx, ∫, Σ, Π, lim, |x|, ^ for powers, / for fractions, parentheses for grouping.
- Convert LaTeX-like tokens to readable math (e.g., \\frac{dy}{dx} -> d/dx, \\sqrt{x} -> sqrt(x)). Remove LaTeX backslashes and $.
- Keep symbols like →, ≈, ≤, ≥ when present; if unclear, use ->.
- If any token is unreadable, write [illegible]; do not invent content.
- If diagrams/figures exist, write [diagram] on its own line; for tables, transcribe cell text line by line.
- Keep the original language. No commentary, no extra instructions.

OUTPUT FORMAT
- Plain text only with original line breaks. No JSON. No markdown. No headings."""
        st.success("✅ OCR prompt force reset!")
        st.rerun()
    
    # Reset buttons
    st.divider()
    if st.button("🔄 Reset All Prompts to Default"):
        st.session_state.grading_prompt = """You are an expert mathematics professor grading multiple questions. Analyze the student's work for ALL questions below.

**QUESTIONS TO GRADE:**
{questions_text}

**GRADING CRITERIA:**
- **IGNORE crossed-out content** - only grade final, uncrossed work
- **Check completion** - penalize incomplete solutions (e.g., "100x-50x" without "=50x")
- **EXCELLENT (90-100%)**: Complete, correct, proper notation, clear reasoning
- **GOOD (70-89%)**: Mostly correct, minor errors, nearly complete
- **FAIR (50-69%)**: Partial understanding, significant errors, incomplete
- **POOR (20-49%)**: Major errors, wrong approach, abandoned work
- **NOT ATTEMPTED (0%)**: No relevant work

**TASK:** Grade each question by finding the relevant work in the images and evaluating according to the criteria above.

**REQUIRED JSON OUTPUT:**
{{
    "question_analysis": [
        {{
            "question_number": <question_number>,
            "status": "<'Excellent'|'Good'|'Fair'|'Poor'|'Not Attempted'>",
            "score": <integer between 0 and max_score>,
            "max_score": <max_score>,
            "feedback": "<brief feedback covering: strengths, errors, completion status, and improvement suggestions>",
            "extracted_work": "<description of work found for this question>",
            "mathematical_quality": "<assessment of rigor and notation>",
            "completion_status": "<'Complete'|'Nearly Complete'|'Partially Complete'|'Incomplete'|'Abandoned'>"
        }}
    ]
}}"""
        st.session_state.summary_prompt = """You are a senior mathematics education specialist with expertise in curriculum development and student assessment. You are creating a comprehensive academic evaluation report for a student's mathematical performance.

**STUDENT'S DETAILED TEST RESULTS:**
{full_analysis}

**CURRICULUM STANDARDS & RUBRIC:**
"{rubric}"

**ACADEMIC EVALUATION REQUIREMENTS:**

**1. OVERALL PERFORMANCE ASSESSMENT:**
Create a comprehensive academic summary (3-4 sentences) that includes:
- **Performance Level**: Overall grade/performance classification
- **Academic Strengths**: Specific mathematical competencies demonstrated
- **Learning Gaps**: Identified areas of conceptual weakness
- **Growth Indicators**: Evidence of mathematical thinking development
- **Strategic Recommendations**: Specific, actionable study and practice suggestions

**2. CONCEPTUAL MASTERY ANALYSIS:**
For each core mathematical concept, provide:
- **Concept Identification**: Clear definition of the mathematical concept
- **Performance Assessment**: Detailed evaluation of understanding level
- **Evidence-Based Analysis**: Specific examples from student work
- **Learning Progression**: Where the student stands in concept mastery
- **Next Steps**: Specific recommendations for concept development

**3. MATHEMATICAL THINKING EVALUATION:**
Assess the student's mathematical reasoning abilities:
- **Problem-Solving Approach**: How they tackle mathematical challenges
- **Logical Reasoning**: Quality of mathematical argumentation
- **Notational Proficiency**: Command of mathematical language
- **Computational Accuracy**: Precision in calculations
- **Conceptual Connections**: Ability to link related mathematical ideas

**4. EDUCATIONAL RECOMMENDATIONS:**
Provide specific, actionable guidance for:
- **Immediate Focus Areas**: What to study next
- **Practice Strategies**: How to improve specific skills
- **Concept Reinforcement**: Ways to strengthen understanding
- **Advanced Preparation**: Preparation for next level mathematics

**REQUIRED JSON OUTPUT:**
{{
    "overall_remarks": "<comprehensive academic summary covering performance level, strengths, gaps, and strategic recommendations>",
    "rubric_analysis": [
        {{
            "concept": "<specific mathematical concept from rubric>",
            "performance": "<'Mastery'|'Proficient'|'Developing'|'Emerging'|'Not Demonstrated'>",
            "evidence": "<detailed evidence from student work with specific examples and mathematical reasoning>",
            "recommendations": "<specific, actionable recommendations for this concept>"
        }}
    ],
    "mathematical_thinking": "<assessment of problem-solving approach, logical reasoning, and mathematical communication>",
    "learning_recommendations": "<comprehensive study and practice recommendations for continued mathematical development>"
}}"""
        st.success("✅ Prompts reset to default! OCR prompts are now configured per question.")
        st.session_state.ocr_prompt = """Extract all text from this image and convert it to human-readable format.

IMPORTANT FORMATTING RULES:
1. Convert all mathematical notation to readable text:
   - Fractions: Write as "numerator/denominator" (e.g., "dy/dx" not "\\frac{dy}{dx}")
   - Square roots: Write as "sqrt(expression)" (e.g., "sqrt(1-x^2)" not "\\sqrt{1-x^2}")
   - Powers: Write as "base^exponent" (e.g., "x^2" not "x^{2}")
   - Logarithms: Write as "log(expression)" (e.g., "log(x)" not "\\log(x)")
   - Trigonometric functions: Write as "sin(x)", "cos(x)", "tan(x)" etc.
   - Greek letters: Write as "alpha", "beta", "theta", "pi" etc.

2. Make the text flow naturally and be easy to read
3. Preserve the mathematical meaning but in plain text format
4. Use standard keyboard symbols instead of LaTeX notation

Return only the converted text content, no explanations."""
        st.success("✅ All prompts reset to default!")
    
    # View saved reports section
    st.divider()
    st.subheader("📁 Saved Reports")
    st.caption("View and download previously saved grading reports")
    
    if os.path.exists("grading_reports"):
        report_files = [f for f in os.listdir("grading_reports") if f.endswith('.json')]
        if report_files:
            selected_report = st.selectbox(
                "Select a report to view:",
                report_files,
                help="Choose a previously saved grading report"
            )
            
            if selected_report:
                try:
                    with open(f"grading_reports/{selected_report}", 'r', encoding='utf-8') as f:
                        report_data = json.load(f)
                    
                    st.json(report_data)
                    
                    # Download button for selected report
                    with open(f"grading_reports/{selected_report}", 'r', encoding='utf-8') as f:
                        file_content = f.read()
                    
                    st.download_button(
                        label="📥 Download This Report",
                        data=file_content,
                        file_name=selected_report,
                        mime="application/json",
                        key=f"sidebar_download_{selected_report}"
                    )
                except Exception as e:
                    st.error(f"Error loading report: {e}")
        else:
            st.info("No saved reports found.")
    else:
        st.info("No reports directory found.")

st.title("💯 Start Grading (Professional Engine)")

# Load OpenAI API key from Streamlit secrets (for deployment) or local config
try:
    # First try Streamlit secrets (for deployment)
    if hasattr(st, 'secrets') and st.secrets:
        OPENAI_API_KEY = st.secrets["openai"]["api_key"]
    else:
        # Fallback to local config files for development
        config_paths = [
            "config.toml",
            "../config.toml",
            ".streamlit/secrets.toml",
            os.path.join(os.path.dirname(__file__), "..", "config.toml"),
            os.path.join(os.getcwd(), "config.toml")
        ]
        
        config = None
        for path in config_paths:
            try:
                config = toml.load(path)
                break
            except FileNotFoundError:
                continue
        
        if config is None:
            st.error(f"🚨 No configuration found! Please set up Streamlit secrets or create config.toml")
            st.stop()
        
        OPENAI_API_KEY = config["openai"]["api_key"]
        
except (KeyError, Exception) as e:
    st.error(f"🚨 Error loading API key: {e}")
    st.stop()


# --- Initialize Session State ---
if "tests" not in st.session_state:
    st.session_state.tests = []
if "students" not in st.session_state:
    st.session_state.students = []
if "submissions" not in st.session_state:
    st.session_state.submissions = {}
if "question_database" not in st.session_state:
    st.session_state.question_database = {}  # Store all questions with their modifications


# --- Helper Functions for Question Management ---
def update_question_in_database(test_title, question_num, new_question_text, max_score):
    """Update a question in the global question database"""
    sanitized_title = test_title.replace(" ", "_").replace("-", "_").replace("(", "").replace(")", "")
    question_key = f"{sanitized_title}_Q{question_num}"
    
    st.session_state.question_database[question_key] = {
        "test_title": test_title,
        "question_num": question_num,
        "question_text": new_question_text,
        "max_score": max_score,
        "last_modified": datetime.now().isoformat()
    }
    
    # Also update the questions list in session state
    questions_key = f"questions_{sanitized_title}"
    if questions_key in st.session_state:
        questions = st.session_state[questions_key]
        for i, (q_num, q_text, q_score) in enumerate(questions):
            if q_num == question_num:
                questions[i] = (question_num, new_question_text, max_score)
                st.session_state[questions_key] = questions
                break

def get_question_from_database(test_title, question_num):
    """Get a question from the global question database"""
    sanitized_title = test_title.replace(" ", "_").replace("-", "_").replace("(", "").replace(")", "")
    question_key = f"{sanitized_title}_Q{question_num}"
    
    if question_key in st.session_state.question_database:
        return st.session_state.question_database[question_key]["question_text"]
    return None

def get_all_questions_for_test(test_title):
    """Get all questions for a test from the database"""
    sanitized_title = test_title.replace(" ", "_").replace("-", "_").replace("(", "").replace(")", "")
    questions = []
    
    # Get questions from database
    for key, value in st.session_state.question_database.items():
        if key.startswith(f"{sanitized_title}_Q"):
            questions.append((value["question_num"], value["question_text"], value["max_score"]))
    
    # Sort by question number (handle both string and integer)
    questions.sort(key=lambda x: int(x[0]) if isinstance(x[0], str) else x[0])
    return questions

def save_updated_results_to_database(test_id, student_id, question_num, new_analysis):
    """Save updated results to a persistent database"""
    if "results_database" not in st.session_state:
        st.session_state.results_database = {}
    
    # Create a unique key for this result
    result_key = f"{test_id}_{student_id}_Q{question_num}"
    
    # Store the updated result with timestamp
    st.session_state.results_database[result_key] = {
        "test_id": test_id,
        "student_id": student_id,
        "question_num": question_num,
        "analysis": new_analysis,
        "updated_at": datetime.now().isoformat(),
        "updated_by": "custom_prompt"  # Could be enhanced to track user
    }
    
    return True

# --- Helper and AI Functions (Unchanged) ---
def extract_text_from_pdf(file):
    try:
        # First try to extract text directly from PDF
        doc = fitz.open(stream=file.read(), filetype="pdf")
        text = "".join(page.get_text() for page in doc)
        doc.close()
        
        # If we got meaningful text, return it
        if text and len(text.strip()) > 10:
            return text
        
        # If no meaningful text, convert PDF to images and use OCR
        file.seek(0)  # Reset file pointer
        images = convert_pdf_to_images(file)
        
        if images:
            # Use our improved OCR on the first page
            extracted_text = extract_text_from_image(images[0])
            return extracted_text
        
        return text  # Return whatever we got
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return None

def convert_pdf_to_images(file):
    images = []
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        for page in doc:
            # Use higher DPI for better math OCR fidelity
            pix = page.get_pixmap(dpi=300)
            img_bytes = pix.tobytes("png")
            image = Image.open(io.BytesIO(img_bytes))
            images.append(image)
        doc.close()
        return images
    except Exception as e:
        st.error(f"Error converting PDF to images: {e}")
        return []

def regenerate_ocr_for_question(question_num, custom_prompt, answer_images, test_id, student_id):
    """Regenerate OCR for a specific question using custom prompt.

    Implementation details:
    - Determines relevant images for the question using weak heuristics; if none match, falls back to all images
    - Sends ALL relevant images in a single API call so the model has page context and can stitch the answer
    - Returns a list of replacements with the combined extracted text tied to each affected image_number
    """
    try:
        # Check if API key is available
        if not OPENAI_API_KEY:
            return None, "OpenAI API key not found"
        
        # Validate inputs
        if not answer_images:
            return None, "No answer images provided"
        # Find images that contain this question (fallback to all if not detectable)
        relevant_images = []
        for answer in answer_images:
            if answer.get('extracted_text'):
                if f"{question_num}." in answer['extracted_text'] or f"{question_num})" in answer['extracted_text']:
                    relevant_images.append(answer)
        # Fallback: if we cannot detect question markers, process all images
        if not relevant_images:
            relevant_images = answer_images
        
        # Use the custom prompt or default
        ocr_prompt = custom_prompt if custom_prompt else st.session_state.ocr_prompt
        
        # Get the original images from session state
        original_images = st.session_state.get(f"original_images_{test_id}_{student_id}", [])
        
        if not original_images:
            return None, "Original images not found in session state"
        
        # Build a single batched call with all relevant images to improve stitching
        import base64
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)

        content_parts = [{"type": "text", "text": ocr_prompt}]
        image_numbers = []
        for answer in relevant_images:
            image_index = answer["image_number"] - 1
            if image_index < len(original_images):
                img_buffer = io.BytesIO()
                original_images[image_index].save(img_buffer, format='PNG')
                img_str = base64.b64encode(img_buffer.getvalue()).decode()
                content_parts.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_str}"}})
                image_numbers.append(answer["image_number"])
                
                response = client.chat.completions.create(
                    model="gpt-4o",
            messages=[{"role": "user", "content": content_parts}],
                    temperature=0.0
                )
                
        combined_text = response.choices[0].message.content.strip()

        # Prepare replacement entries for each affected image number so callers can merge in-place
        new_extractions = []
        for num in image_numbers:
            new_extractions.append({
                "image_number": num,
                "extracted_text": combined_text,
                "character_count": len(combined_text)
            })

        return new_extractions, "OCR regenerated successfully (batched)"
    except Exception as e:
        error_msg = f"Error regenerating OCR: {str(e)}"
        if st.session_state.get('debug_mode', False):
            st.error(f"Debug: {error_msg}")
            import traceback
            st.write(f"Debug: Full traceback: {traceback.format_exc()}")
        return None, error_msg

def regenerate_analysis_for_question(question_num, custom_prompt, question_text, student_answer, rubric, test_id, student_id, extracted_answers=None):
    """Regenerate analysis for a specific question using custom prompt"""
    try:
        # Debug information
        if st.session_state.get('debug_mode', False):
            st.write(f"Debug: regenerate_analysis_for_question called with:")
            st.write(f"- question_num: {question_num}")
            st.write(f"- question_text: {question_text[:100]}...")
            st.write(f"- student_answer: {student_answer[:100] if student_answer else 'None'}...")
            st.write(f"- extracted_answers: {len(extracted_answers) if extracted_answers else 0} items")
        
        # Check if API key is available
        if not OPENAI_API_KEY:
            return None, "OpenAI API key not found"
        
        # Validate inputs
        if not question_text:
            return None, "Missing question text"
        
        # If no specific student answer found, use all extracted text
        if not student_answer or student_answer.strip() == "":
            # Get all extracted text as fallback
            all_text = ""
            if extracted_answers:
                for answer in extracted_answers:
                    if answer.get('extracted_text'):
                        all_text += f"\n{answer['extracted_text']}\n"
                student_answer = all_text.strip()
            
            if not student_answer:
                # Debug information
                if st.session_state.get('debug_mode', False):
                    st.error(f"Debug: No student work found. extracted_answers: {extracted_answers}")
                return None, "No student work found in any images. Please check if answer images were uploaded and OCR extraction was successful."
        # Create a focused prompt for this specific question
        if custom_prompt:
            # Add JSON requirement to custom prompt
            if "json" not in custom_prompt.lower():
                analysis_prompt = f"""{custom_prompt}

Please return your analysis in JSON format:
{{
    "score": <score>,
    "max_score": <max_score>,
    "status": "<status>",
    "feedback": "<detailed feedback>",
    "extracted_work": "<work summary>",
    "mathematical_quality": "<quality assessment>",
    "completion_status": "<completion status>"
}}"""
            else:
                analysis_prompt = custom_prompt
        else:
            analysis_prompt = f"""Analyze the student's work for Question {question_num}.

Question: {question_text}
Student's Answer: {student_answer}

Provide a detailed analysis including:
1. Score (0 to max score)
2. Status (Excellent/Good/Fair/Poor/Not Attempted)
3. Detailed feedback
4. Work summary
5. Mathematical quality assessment
6. Completion status

Return your analysis in JSON format:
{{
    "score": <score>,
    "max_score": <max_score>,
    "status": "<status>",
    "feedback": "<detailed feedback>",
    "extracted_work": "<work summary>",
    "mathematical_quality": "<quality assessment>",
    "completion_status": "<completion status>"
}}"""

        # Debug: Print the prompt being sent
        if st.session_state.get('debug_mode', False):
            st.write(f"Debug: Sending analysis prompt to GPT:")
            st.text_area("Analysis Prompt", analysis_prompt, height=200, disabled=True)

        # Call GPT for analysis
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "user", "content": analysis_prompt}
            ],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        
        # Debug: Print the raw response
        if st.session_state.get('debug_mode', False):
            st.write(f"Debug: GPT Response received:")
            st.write(f"Debug: Response content: {response.choices[0].message.content}")
        
        result = json.loads(response.choices[0].message.content)
        
        # Debug: Print the parsed result
        if st.session_state.get('debug_mode', False):
            st.write(f"Debug: Parsed JSON result: {result}")
        
        return result, "Analysis regenerated successfully"
    except json.JSONDecodeError as e:
        error_msg = f"JSON parsing error: {str(e)}. Raw response: {response.choices[0].message.content if 'response' in locals() else 'No response'}"
        if st.session_state.get('debug_mode', False):
            st.error(f"Debug: {error_msg}")
        return None, error_msg
    except Exception as e:
        error_msg = f"Error regenerating analysis: {str(e)}"
        if st.session_state.get('debug_mode', False):
            st.error(f"Debug: {error_msg}")
            import traceback
            st.write(f"Debug: Full traceback: {traceback.format_exc()}")
        return None, error_msg

def extract_question_answer(extracted_answers, question_num):
    """Extract the answer span strictly from the start of the given answer number
    up to the start of the next answer number. Uses line-anchored patterns to
    avoid accidental mid-line matches and supports common formats like
    "1.", "1)", "Question 1", "Q1".
    """
    if not extracted_answers:
        return ""
    
    # Combine all extracted text and normalize newlines/whitespace
    all_text_parts = []
    for answer in extracted_answers:
        text = answer.get("extracted_text") or ""
        if not isinstance(text, str):
            continue
        all_text_parts.append(text)
    all_text = "\n".join(all_text_parts)
    all_text = all_text.replace("\r\n", "\n").replace("\r", "\n")

    # Helper to build a robust, line-anchored pattern for a question number
    def build_line_anchored_pattern(n: int) -> str:
        # Matches at the beginning of a line: optional 'Question' or 'Q', spaces, the number,
        # followed by optional punctuation and a space. Examples:
        # 1. , 1) , 1: , Q1 , Question 1 , Q.1 , (1) etc.
        return (
            rf"(?mi)^(?:\s*(?:question\s*)?q?\s*{n}\s*(?:[\.)\]:-])?\s+)"
        )

    # Find start anchored to a line
    start_pattern = build_line_anchored_pattern(question_num)
    start_match = re.search(start_pattern, all_text)
    if not start_match:
        # Fallback: previous relaxed patterns if strict anchor fails
        relaxed_patterns = [
            rf"{question_num}\.\s*",
            rf"{question_num}\)\s*",
            rf"Question\s*{question_num}\s*",
            rf"Q{question_num}\s*",
            rf"\b{question_num}\b",
        ]
        for pat in relaxed_patterns:
            m = re.search(pat, all_text, re.IGNORECASE)
            if m:
                start_match = m
            break
    
    if not start_match:
        # If still not found, return the full text as a last resort
        return all_text.strip()
    
    start_pos = start_match.start()

    # Find the next question start anchored to a line
    next_pattern = build_line_anchored_pattern(question_num + 1)
    next_match = re.search(next_pattern, all_text[start_pos:])
    if next_match:
        end_pos = start_pos + next_match.start()
    else:
        # If not found, search relaxed next markers
        relaxed_next_patterns = [
            rf"{question_num + 1}\.\s*",
            rf"{question_num + 1}\)\s*",
            rf"Question\s*{question_num + 1}\s*",
            rf"Q{question_num + 1}\s*",
            rf"\b{question_num + 1}\b",
        ]
        end_pos = len(all_text)
        for pat in relaxed_next_patterns:
            m = re.search(pat, all_text[start_pos:], re.IGNORECASE)
            if m:
                end_pos = start_pos + m.start()
            break
    
    # Slice and clean
    span = all_text[start_pos:end_pos]
    # Trim leading/trailing blank lines but keep internal structure
    lines = [ln.rstrip() for ln in span.split("\n")]
    # Drop obvious scanning noise lines that are empty after stripping
    cleaned_lines = [ln for ln in lines if ln.strip() != ""]
    final_answer = "\n".join(cleaned_lines).strip()

    if st.session_state.get("debug_mode", False):
        st.write(f"Debug: start_pos={start_pos}, end_pos={end_pos}, length={len(final_answer)}")
        st.text_area(f"Debug: Extracted Answer Q{question_num}", final_answer, height=120, disabled=True)
    
    return final_answer

def get_current_answer_for_question(submission: dict, question_num: int) -> str:
    """Return the best-current answer text for a question.

    Preference order:
    1) Manually edited override stored in submission['answer_overrides']
    2) Text extracted from images in submission['extracted_answers']
    """
    if not submission:
        return ""

    overrides = submission.get('answer_overrides') or {}
    # Support both str and int keys gracefully
    override_val = overrides.get(str(question_num)) or overrides.get(question_num)
    if isinstance(override_val, str) and override_val.strip():
        return override_val.strip()

    extracted_answers = submission.get('extracted_answers', [])
    return extract_question_answer(extracted_answers, question_num)

def extract_text_from_image(image, custom_prompt=None):
    """Extract text from image using GPT-4o OCR with human-readable formatting"""
    try:
        # Convert image to base64 for OpenAI API
        import base64
        img_buffer = io.BytesIO()
        image.save(img_buffer, format='PNG')
        img_str = base64.b64encode(img_buffer.getvalue()).decode()
        
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        
        # Use custom prompt if provided, otherwise use the session state OCR prompt
        ocr_prompt = custom_prompt if custom_prompt else st.session_state.ocr_prompt
        
        # Show the OCR prompt only in debug mode to avoid UI spam during loops
        if st.session_state.get('debug_mode', False):
            st.info(f"🔍 Using OCR prompt: {ocr_prompt[:100]}...")
            st.text_area("OCR Prompt", ocr_prompt, height=120, disabled=True)
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "user", "content": [
                    {"type": "text", "text": ocr_prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_str}"}}
                ]}
            ],
            temperature=0.0
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"GPT-4o OCR failed: {str(e)}"

def extract_topics_from_rubric(rubric_text: str) -> list:
    """Parse rubric text to produce a list of topics.

    Splits on lines and commas; removes numbering; drops duplicates and tiny tokens.
    """
    if not rubric_text:
        return []
    candidates = []
    for raw_line in rubric_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        line = re.sub(r"^\s*(\d+\.|\d+\)|-\s+|•\s+)", "", line)
        parts = [p.strip() for p in re.split(r",|;|/", line) if p.strip()]
        candidates.extend(parts if parts else [line])
    seen = set()
    topics = []
    for c in candidates:
        norm = re.sub(r"\s+", " ", c).strip()
        if len(norm) < 3:
            continue
        key = norm.lower()
        if key not in seen:
            seen.add(key)
            topics.append(norm)
    return topics

def analyze_topic_coverage(topics: list, questions: list, extracted_answers: list) -> list:
    """Return coverage per topic using simple keyword presence checks."""
    question_text_blob = "\n".join([(q[1] if len(q) >= 2 else str(q)) for q in (questions or [])]).lower()
    answers_blob = "\n".join([a.get('extracted_text', '') for a in (extracted_answers or [])]).lower()
    results = []
    for topic in topics:
        t = topic.lower()
        in_questions = t in question_text_blob
        in_answers = t in answers_blob
        if in_questions and in_answers:
            status = "✅ Covered"
        elif in_questions and not in_answers:
            status = "⚠️ Asked but not demonstrated"
        elif (not in_questions) and in_answers:
            status = "ℹ️ Demonstrated (not explicitly asked)"
        else:
            status = "❌ Missing"
        results.append({
            "concept": topic,
            "in_questions": in_questions,
            "in_answers": in_answers,
            "status": status,
        })
    return results

def extract_rubric_table_from_text(rubric_text: str) -> list:
    """Extract a structured rubric table from free-form rubric text.

    Returns a list of rows with the following keys:
    - "Concept No.": int
    - "Concept": str
    - "Example": str
    - "Status": str
    """
    if not rubric_text or len(rubric_text.strip()) == 0:
        return []
    try:
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompt = f"""
You will receive RUBRIC TEXT that may include a table or bullet list of concepts and examples.
Extract it as a clean JSON table with columns: Concept No., Concept, Example, Status.
Rules:
- Number Concept No. starting from 1 in the order found
- If Example is missing for a row, set it to ""
- Set Status to "" for all rows (left blank for evaluation later)
Return ONLY a JSON object of the shape:
{{
  "rows": [{{"Concept No.": 1, "Concept": "...", "Example": "...", "Status": ""}}, ...]
}}

RUBRIC TEXT:\n{rubric_text}
"""
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        data = json.loads(response.choices[0].message.content)
        rows = data.get("rows", [])
        # Basic sanitation
        normalized = []
        for idx, r in enumerate(rows, start=1):
            normalized.append({
                "Concept No.": int(r.get("Concept No.", idx)),
                "Concept": str(r.get("Concept", "")).strip(),
                "Example": str(r.get("Example", "")).strip(),
                "Status": str(r.get("Status", "")).strip(),
            })
        return normalized
    except Exception:
        # Heuristic fallback: split by commas and newlines into concept list
        topics = [t.strip() for t in re.split(r",|\n", rubric_text) if t.strip()]
        return [{"Concept No.": i + 1, "Concept": t, "Example": "", "Status": ""} for i, t in enumerate(topics)]

def parse_rubric_text_llm(rubric_text: str) -> list:
    """Parse rubric from plain text using a single LLM call. Preserves math.

    Returns list of rows with Concept No., Concept, Example, Status.
    """
    if not rubric_text or len(rubric_text.strip()) == 0:
        return []
    try:
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompt = (
            "You are given RUBRIC TEXT copy-pasted from a PDF containing a rubric table.\n"
            "Reconstruct a JSON table with columns: Concept No., Concept, Example, Status.\n\n"
            "Rules:\n"
            "- Parse rows in order. If numbering is missing, infer sequential Concept No. starting at 1.\n"
            "- In Example, TRANSCRIBE MATH FAITHFULLY: keep ∫, √, d/dx, Σ, Π, |x|; convert LaTeX-like tokens (\\frac{dy}{dx} -> d/dx; \\sqrt{x} -> √x; x^{2} -> x^2; x_{0} -> x_0).\n"
            "- If multiple snippets exist (i), (ii), (iii), include ALL in one cell separated by '; '.\n"
            "- Status should be empty string for all rows.\n"
            "Return ONLY JSON: {\"rows\": [{\"Concept No.\": 1, \"Concept\": \"…\", \"Example\": \"…\", \"Status\": \"\"}, …]}.\n\n"
            f"RUBRIC TEXT\n{rubric_text}"
        )
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        data = json.loads(resp.choices[0].message.content)
        rows = data.get("rows", [])
        normalized = []
        for idx, r in enumerate(rows, start=1):
            normalized.append({
                "Concept No.": int(r.get("Concept No.", idx)),
                "Concept": str(r.get("Concept", "")).strip(),
                "Example": str(r.get("Example", "")).strip(),
                "Status": str(r.get("Status", "")).strip(),
            })
        return normalized
    except Exception:
        return extract_rubric_table_from_text(rubric_text)

def fill_missing_examples_from_text(rows: list, rubric_text: str) -> list:
    """For rows with empty Example, ask the LLM to locate and transcribe the Example
    cell from the rubric text. Updates and returns rows."""
    if not rows:
        return []
    if not rubric_text:
        return rows
    try:
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        for i, row in enumerate(rows):
            if str(row.get("Example", "")).strip():
                continue
            concept_no = row.get("Concept No.") or (i + 1)
            concept_name = str(row.get("Concept", "")).strip()
            ask = (
                "From the following rubric text, find the table row that matches the given concept number/name,\n"
                "and output EXACTLY the Example cell text for that row.\n\n"
                f"Target row: Concept No. = {concept_no}; Concept contains: '{concept_name[:120]}'\n\n"
                "Rubric Text (may include the entire table as text):\n" + rubric_text + "\n\n"
                "Rules: preserve math: ∫, √, Σ, Π, d/dx, |x|, powers ^, fractions /. Convert LaTeX-like tokens to readable math.\n"
                "If multiple items exist (i), (ii), keep them separated by '; '.\n"
                "Return plain text only; no quotes, no labels."
            )
            resp = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": ask}],
                temperature=0.0,
            )
            text = (resp.choices[0].message.content or "").strip().strip('"')
            if text:
                rows[i]["Example"] = text
        return rows
    except Exception:
        return rows
def extract_rubric_rows_from_images(images: list) -> list:
    """Extract a full rubric table from one or more rubric images using GPT-4o.

    Attempts to capture ALL columns present in the table across pages and returns
    a list of row dicts. Known columns like 'Concept No.', 'Concept', 'Example',
    'Status' are normalized when possible, but any additional columns found are
    preserved as extra keys on each row.
    """
    if not images:
        return []
    try:
        import base64
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)

        instruction = (
            "You are given one or more images that together contain a rubric table.\n"
            "Extract the ENTIRE table accurately, including ALL columns and ALL rows,\n"
            "merging content across images if the table spans multiple pages.\n\n"
            "Return STRICT JSON with these fields only:\n"
            "{\n  \"columns\": [\"<header1>\", \"<header2>\", ...],\n  \"rows\": [ {\"<header1>\": \"cell\", ...}, ... ]\n}\n\n"
            "Rules:\n"
            "- Detect the header row; use the exact visible header text as column names.\n"
            "- Preserve cell text faithfully (no summarization).\n"
            "- For the 'Example' column, TRANSCRIBE MATH PRECISELY and COMPLETELY. Keep symbols: ∫, √, π, θ, Σ, Π, |x|, d/dx, lim; trig/log names.\n"
            "  Convert LaTeX-like tokens to readable math: \\\\frac{dy}{dx} -> d/dx; \\\\sqrt{x} -> √x; x^{2} -> x^2; x_{0} -> x_0.\n"
            "  If multiple snippets are listed (i), (ii), (iii), include ALL of them separated by '; '.\n"
            "- If a cell is empty, use an empty string.\n"
            "- If a cell spans multiple lines, join with a single space.\n"
            "- If a header is missing but implied, infer a short header like 'Column X'.\n"
            "- If the table does not have explicit numbering, do not invent numbers; leave numbering column blank.\n"
            "- Do NOT add commentary. Output JSON object only."
        )

        content_parts = [{"type": "text", "text": instruction}]
        for img in images:
            buf = io.BytesIO()
            img.save(buf, format='PNG')
            img_b64 = base64.b64encode(buf.getvalue()).decode()
            content_parts.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}})

        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": content_parts}],
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        data = json.loads(resp.choices[0].message.content)

        # Accept either {columns, rows} or legacy {rows}
        rows = []
        if isinstance(data, dict):
            if "rows" in data and isinstance(data.get("rows"), list):
                rows = data["rows"]

        # Normalize best-effort for known columns while preserving extras
        def get_first_present_key(d: dict, candidates: list, default: str = "") -> str:
            for k in candidates:
                if k in d and str(d.get(k)).strip() != "":
                    return str(d.get(k))
            return default

        normalized_rows = []
        for idx, r in enumerate(rows, start=1):
            # Copy all original cells
            row_out = {k: ("" if r.get(k) is None else str(r.get(k)).strip()) for k in r.keys()}
            # Inject normalized standard columns
            concept_no = get_first_present_key(r, [
                "Concept No.", "No.", "No", "S.No", "S. No.", "#", "ID"
            ])
            if concept_no == "":
                concept_no = str(idx)
            row_out["Concept No."] = int(str(concept_no).split()[0].strip(".:)")) if str(concept_no).split() else idx

            row_out["Concept"] = get_first_present_key(r, [
                "Concept", "Topic", "Outcome", "Learning Outcome", "Skill", "Criteria"
            ])
            row_out["Example"] = get_first_present_key(r, [
                "Example", "Illustration", "Sample", "Formula", "Notes"
            ])
            # Keep Status blank unless present
            row_out["Status"] = get_first_present_key(r, [
                "Status", "Remark", "Notes/Status"
            ], default="")
            normalized_rows.append(row_out)

        return normalized_rows
    except Exception:
        return []


def _text_has_math_indicators(text: str) -> bool:
    if not isinstance(text, str):
        return False
    lowered = text.lower()
    return any(tok in text for tok in ["∫", "√", "Σ", "Π", "^", "|", "→"]) or \
        any(fn in lowered for fn in ["sin", "cos", "tan", "log", "ln", "lim", "dx", "dy/dx", "sec", "cosec", "cot"]) or \
        ("/" in text and any(ch.isdigit() for ch in text))


def refine_rubric_rows_with_cell_ocr(images: list, rows: list) -> list:
    """Second pass: for each row, if the Example cell looks empty or too generic,
    run a targeted OCR over the rubric images to locate and transcribe the Example cell
    corresponding to that row. Updates rows in-place and returns them.
    """
    if not images or not rows:
        return rows or []
    try:
        import base64
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)

        # Pre-encode images once
        image_parts = []
        for img in images:
            buf = io.BytesIO()
            img.save(buf, format='PNG')
            img_b64 = base64.b64encode(buf.getvalue()).decode()
            image_parts.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}})

        for idx, row in enumerate(rows):
            concept_no = row.get("Concept No.") or idx + 1
            concept_name = str(row.get("Concept", "")).strip()
            example_text = str(row.get("Example", "")).strip()

            # Heuristic: refine if missing or no clear math indicators
            if len(example_text) >= 10 and _text_has_math_indicators(example_text):
                continue

            instruction = (
                "You are given images of a rubric table. Find the row that matches the given Concept number/name,\n"
                "then transcribe EXACTLY the text from its 'Example' cell.\n\n"
                f"Target row: Concept No. = {concept_no}; Concept/Title contains: '{concept_name[:120]}'\n\n"
                "Rules:\n"
                "- Output only the exact Example cell text for that row. No headings, no extra commentary.\n"
                "- Preserve math faithfully: keep ∫, √, Σ, Π, |x|, d/dx, lim, etc.\n"
                "- Convert LaTeX-like tokens to readable math: \\frac{dy}{dx} -> d/dx; \\sqrt{x} -> √x; x^{2} -> x^2; x_{0} -> x_0.\n"
                "- If multiple short items exist (i), (ii), (iii), keep them as '(i) …; (ii) …; (iii) …' in one line.\n"
                "- If unreadable, write [illegible]."
            )

            content = [{"type": "text", "text": instruction}] + image_parts
            resp = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": content}],
                temperature=0.0,
            )
            extracted = (resp.choices[0].message.content or "").strip()
            if extracted:
                # Avoid model adding quotes or labels
                cleaned = extracted.strip().strip('"')
                rows[idx]["Example"] = cleaned

        return rows
    except Exception:
        return rows


def merge_rubric_rows_by_concept(base_rows: list, image_rows: list) -> list:
    """Merge rows extracted from images into base rows (usually text-parsed).
    Uses Concept No. primarily, then falls back to normalized Concept text.
    Replaces Example/Status if provided by image rows; preserves extra columns.
    """
    if not base_rows:
        return image_rows or []
    if not image_rows:
        return base_rows

    def norm_name(v: str) -> str:
        return re.sub(r"\s+", " ", (v or "").strip().lower())

    # Build lookup from image rows
    by_no = {int(r.get("Concept No.", 0)): r for r in image_rows if str(r.get("Concept No."))}
    by_name = {norm_name(r.get("Concept", "")): r for r in image_rows}

    merged: list = []
    for r in base_rows:
        out = dict(r)
        candidate = by_no.get(int(r.get("Concept No.", 0))) or by_name.get(norm_name(r.get("Concept", "")))
        if candidate:
            # Replace Example/Status if present in image row
            ex = str(candidate.get("Example", "")).strip()
            if ex:
                out["Example"] = ex
            st_val = str(candidate.get("Status", "")).strip()
            if st_val:
                out["Status"] = st_val
            # Preserve any extra columns from image row
            for k, v in candidate.items():
                if k not in out:
                    out[k] = v
        merged.append(out)
    return merged

def _status_to_score(status: str) -> int:
    mapping = {
        'excellent': 5,
        'good': 4,
        'fair': 3,
        'poor': 2,
        'not attempted': 1,
        'unknown': 0,
    }
    return mapping.get(str(status or '').strip().lower(), 0)

def _score_to_status(score: float) -> str:
    if score >= 4.5:
        return 'Excellent'
    if score >= 3.5:
        return 'Good'
    if score >= 2.5:
        return 'Fair'
    if score >= 1.5:
        return 'Poor'
    if score > 0:
        return 'Not Attempted'
    return 'Unknown'

def build_completed_rubric_table(rubric_rows: list, questions: list, question_analysis: list) -> list:
    """Produce a completed rubric table with asked question numbers and implementation quality.

    - asked question numbers are detected by simple substring match of topic in question text
    - implementation is the average of the mapped statuses for the asked questions
    """
    # Prepare question texts map
    qnum_to_text = {}
    if questions:
        for q in questions:
            try:
                if len(q) == 3:
                    q_num, q_text, _ = int(q[0]), q[1].strip(), int(q[2])
                elif len(q) == 2:
                    q_num, q_text = int(q[0]), q[1].strip()
                else:
                    continue
                qnum_to_text[q_num] = q_text
            except Exception:
                continue

    # Prepare analysis map
    qnum_to_status = {}
    if question_analysis:
        for qa in question_analysis:
            qn = qa.get('question_number')
            status = qa.get('status', 'Unknown')
            if isinstance(qn, int):
                qnum_to_status.setdefault(qn, []).append(status)

    completed_rows = []
    for row in rubric_rows or []:
        concept = (row.get('Concept') or '').strip()
        if not concept:
            continue
        matched_qnums = []
        concept_lower = concept.lower()
        for qn, qtext in qnum_to_text.items():
            try:
                if concept_lower in (qtext or '').lower():
                    matched_qnums.append(qn)
            except Exception:
                continue
        matched_qnums_sorted = sorted(set(matched_qnums))
        asked_str = ", ".join(str(n) for n in matched_qnums_sorted) if matched_qnums_sorted else "—"

        # Aggregate implementation quality
        if matched_qnums_sorted:
            scores = []
            for qn in matched_qnums_sorted:
                statuses = qnum_to_status.get(qn, [])
                if statuses:
                    scores.extend([_status_to_score(s) for s in statuses])
            implementation = _score_to_status(sum(scores) / len(scores)) if scores else 'Unknown'
        else:
            implementation = 'Not Asked'

        completed = dict(row)
        completed["Asked in Questions"] = asked_str
        completed["Implementation"] = implementation
        completed_rows.append(completed)

    return completed_rows

def generate_rubric_summary_table(rubric_table_rows: list) -> list:
    """Ask the model for a short, teacher-friendly summary table.

    Output columns: Concept, Summary, Action.
    """
    if not rubric_table_rows:
        return []
    try:
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompt = {
            "instruction": "Convert the rubric evaluation rows into a concise teacher-facing summary table.",
            "columns": ["Concept", "Summary", "Action"],
            "rows": rubric_table_rows,
        }
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": json.dumps(prompt)}],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        data = json.loads(response.choices[0].message.content)
        return data.get("rows", [])
    except Exception:
        # Fallback: generate a naive summary
        rows = []
        for r in rubric_table_rows:
            rows.append({
                "Concept": r.get("Concept", ""),
                "Summary": f"Asked in Q: {r.get('Asked in Questions', '—')}, Implementation: {r.get('Implementation', 'Unknown')}",
                "Action": "Focus practice on weak/unknown concepts."
            })
        return rows

def save_extracted_content(test_title, student_name, question_text, rubric, answer_images, question_analysis, final_score, summary_data=None):
    """Save all extracted content to a comprehensive DOC file"""
    try:
        # Create reports directory if it doesn't exist
        os.makedirs("grading_reports", exist_ok=True)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_filename = f"grading_reports/{test_title}_{student_name}_{timestamp}.docx"
        
        # Extract text from all answer images
        extracted_answers = []
        for i, img in enumerate(answer_images):
            extracted_text = extract_text_from_image(img)
            extracted_answers.append({
                "image_number": i + 1,
                "extracted_text": extracted_text,
                "character_count": len(extracted_text) if extracted_text else 0
            })
        
        # Create DOC document
        doc = Document()
        
        # Title
        title = doc.add_heading(f'Grading Report: {test_title}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_heading('Report Information', level=1)
        doc.add_paragraph(f'Student: {student_name}')
        doc.add_paragraph(f'Test: {test_title}')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph(f'Total Score: {final_score}/100')
        doc.add_paragraph(f'Total Questions: {len(question_analysis)}')
        
        # Question Paper
        doc.add_heading('Question Paper', level=1)
        if question_text:
            # Format question text properly while preserving structure
            lines = question_text.split('\n')
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if it's a section header
                if line.upper().startswith('SECTION'):
                    current_section = line
                    doc.add_heading(line, level=2)
                
                # Check if it's a question number (1), 2), etc.)
                elif re.match(r'^\d+[\)\.]', line):
                    # Extract question number and text
                    question_match = re.match(r'^(\d+)[\)\.]\s*(.*)', line)
                    if question_match:
                        q_num = question_match.group(1)
                        q_text = question_match.group(2)
                        
                        # Format question with proper spacing
                        question_heading = f"{q_num}) {q_text}"
                        doc.add_heading(question_heading, level=3)
                    else:
                        doc.add_heading(line, level=3)
                
                # Check if it's header information
                elif any(keyword in line.upper() for keyword in ['CLASS', 'TIME', 'MARKS', 'CAS CLASSES', 'DIFFERENTIATION TEST']):
                    # Format header information
                    if 'CAS CLASSES' in line.upper():
                        # Split CAS CLASSES line into multiple lines for better formatting
                        parts = line.split('CAS CLASSES')
                        if len(parts) > 1:
                            doc.add_paragraph('CAS CLASSES', style='Heading 4')
                            doc.add_paragraph(parts[1].strip())
                        else:
                            doc.add_paragraph(line, style='Heading 4')
                    elif 'CLASS 12' in line.upper():
                        doc.add_heading(line, level=4)
                    elif 'TIME ALLOWED' in line.upper() or 'MAXIMUM MARKS' in line.upper():
                        doc.add_paragraph(line, style='Heading 4')
                    else:
                        doc.add_paragraph(line, style='Heading 4')
                
                # Regular question text (continuation of questions)
                else:
                    # Check if this is continuation of a question
                    if current_section and not line.startswith('[') and not line.startswith('Time') and not line.startswith('Maximum'):
                        doc.add_paragraph(line)
                    elif line.startswith('[') and line.endswith(']'):
                        # This is a marks indicator, add it to the previous paragraph
                        continue
                    else:
                        doc.add_paragraph(line)
        else:
            doc.add_paragraph("No question text available")
        
        # Rubric omitted in report per configuration
        
        # Extracted Answers
        doc.add_heading('Student Answers (Extracted from Images)', level=1)
        for answer in extracted_answers:
            doc.add_heading(f'Image {answer["image_number"]}', level=2)
            if answer.get('extracted_text'):
                # Clean and format extracted text
                clean_text = answer['extracted_text'].replace('\n\n', '\n').strip()
                if clean_text:
                    doc.add_paragraph(clean_text)
                else:
                    doc.add_paragraph("No text extracted from this image")
            else:
                doc.add_paragraph("No text extracted from this image")
        
        # Question Analysis
        doc.add_heading('Question-by-Question Analysis', level=1)
        for qa in question_analysis:
            doc.add_heading(f'Question {qa.get("question_number", "N/A")}', level=2)
            doc.add_paragraph(f'Score: {qa.get("score", 0)}/{qa.get("max_score", 0)}')
            doc.add_paragraph(f'Status: {qa.get("status", "Unknown")}')
            
            # Clean and format feedback
            feedback = qa.get("feedback", "No feedback available")
            if feedback:
                clean_feedback = feedback.replace('\n', ' ').strip()
                doc.add_paragraph(f'Feedback: {clean_feedback}')
            
            if qa.get('extracted_work'):
                clean_work = qa.get('extracted_work').replace('\n', ' ').strip()
                doc.add_paragraph(f'Work Found: {clean_work}')
            
            if qa.get('mathematical_quality'):
                clean_quality = qa.get('mathematical_quality').replace('\n', ' ').strip()
                doc.add_paragraph(f'Mathematical Quality: {clean_quality}')
            
            if qa.get('completion_status'):
                doc.add_paragraph(f'Completion Status: {qa.get("completion_status")}')
        
        # Overall Summary
        doc.add_heading('Overall Summary', level=1)
        
        # Check for summary data passed as parameter
        summary_found = False
        if summary_data:
            if summary_data.get('overall_remarks'):
                clean_remarks = summary_data.get('overall_remarks').replace('\n', ' ').strip()
                doc.add_paragraph(clean_remarks)
                summary_found = True
            
            if summary_data.get('mathematical_thinking'):
                doc.add_heading('Mathematical Thinking Assessment', level=2)
                clean_thinking = summary_data.get('mathematical_thinking').replace('\n', ' ').strip()
                doc.add_paragraph(clean_thinking)
                summary_found = True
            
            if summary_data.get('learning_recommendations'):
                doc.add_heading('Learning Recommendations', level=2)
                clean_recommendations = summary_data.get('learning_recommendations').replace('\n', ' ').strip()
                doc.add_paragraph(clean_recommendations)
                summary_found = True
        
        # If no summary found, create a basic summary from scores
        if not summary_found:
            total_score = sum(qa.get('score', 0) for qa in question_analysis)
            max_score = sum(qa.get('max_score', 0) for qa in question_analysis)
            percentage = (total_score / max_score * 100) if max_score > 0 else 0
            
            # Count performance levels
            excellent_count = sum(1 for qa in question_analysis if qa.get('status') == 'Excellent')
            good_count = sum(1 for qa in question_analysis if qa.get('status') == 'Good')
            fair_count = sum(1 for qa in question_analysis if qa.get('status') == 'Fair')
            poor_count = sum(1 for qa in question_analysis if qa.get('status') == 'Poor')
            not_attempted = sum(1 for qa in question_analysis if qa.get('status') == 'Not Attempted')
            
            # Generate summary
            summary_text = f"Student achieved a total score of {total_score}/{max_score} ({percentage:.1f}%). "
            summary_text += f"Performance breakdown: {excellent_count} excellent, {good_count} good, {fair_count} fair, {poor_count} poor, and {not_attempted} not attempted questions. "
            
            if percentage >= 90:
                summary_text += "Overall performance is excellent with strong understanding of the concepts."
            elif percentage >= 80:
                summary_text += "Overall performance is very good with minor areas for improvement."
            elif percentage >= 70:
                summary_text += "Overall performance is good with some areas needing attention."
            elif percentage >= 60:
                summary_text += "Overall performance is satisfactory but requires significant improvement."
            else:
                summary_text += "Overall performance needs substantial improvement and additional practice."
            
            doc.add_paragraph(summary_text)
            
            # Add recommendations based on performance
            doc.add_heading('Recommendations', level=2)
            if poor_count > 0 or not_attempted > 0:
                doc.add_paragraph("Focus on reviewing fundamental concepts and practicing basic differentiation techniques.")
            if fair_count > 0:
                doc.add_paragraph("Work on improving accuracy and understanding of intermediate-level problems.")
            if excellent_count < len(question_analysis) * 0.5:
                doc.add_paragraph("Continue practicing advanced problems to strengthen overall mathematical skills.")
        
        # Save DOC file
        doc.save(doc_filename)
        
        return doc_filename
    except Exception as e:
        st.error(f"Error saving extracted content: {e}")
        return None

def process_uploaded_files(uploaded_files):
    """Process uploaded files (PDFs and images) and return a list of PIL Images"""
    images = []
    
    for uploaded_file in uploaded_files:
        try:
            file_type = uploaded_file.type.lower()
            
            if file_type == "application/pdf":
                # Handle PDF files
                pdf_images = convert_pdf_to_images(uploaded_file)
                images.extend(pdf_images)
            elif file_type.startswith("image/"):
                # Handle image files (jpg, png, jpeg, etc.)
                image = Image.open(uploaded_file)
                images.append(image)
            else:
                st.warning(f"Unsupported file type: {file_type}. Please upload PDF or image files.")
                
        except Exception as e:
            st.error(f"Error processing file {uploaded_file.name}: {e}")
    
    return images

def _grade_all_questions_with_gpt4o(questions, answer_images):
    # Create a concise questions list
    questions_text = ""
    for i, q in enumerate(questions):
        # Handle tuples with different lengths (with or without max score)
        if len(q) == 3:
            q_num, q_text, q_max_score = int(q[0]), q[1].strip(), int(q[2])
        elif len(q) == 2:
            q_num, q_text = int(q[0]), q[1].strip()
            q_max_score = 10  # Default max score
        else:
            continue  # Skip invalid questions
        # Use full question text to preserve modifications
        questions_text += f"Q{q_num} ({q_max_score}pts): {q_text}\n"
    
    # Use the session state prompt with the questions text
    prompt = st.session_state.grading_prompt.format(questions_text=questions_text)
    
    # Debug: Show the actual prompt being used (optional)
    if st.session_state.get('debug_mode', False):
        st.subheader("🔍 Debug: Grading Prompt Being Used")
        st.text_area("Grading Prompt", prompt, height=300, disabled=True)
    
    try:
        # Convert images to base64 for OpenAI API
        import base64
        image_contents = []
        for img in answer_images:
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            img_str = base64.b64encode(img_buffer.getvalue()).decode()
            image_contents.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/png;base64,{img_str}"
                }
            })
        
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "user", "content": [
                    {"type": "text", "text": prompt},
                    *image_contents
                ]}
            ],
            temperature=0.0,
            response_format={"type": "json_object"}
        )
        result = json.loads(response.choices[0].message.content)
        return result.get("question_analysis", [])
    except Exception as e:
        # Return error results for all questions
        error_results = []
        for q in questions:
            # Handle tuples with different lengths (with or without max score)
            if len(q) == 3:
                q_num, q_text, q_max_score = int(q[0]), q[1].strip(), int(q[2])
            elif len(q) == 2:
                q_num, q_text = int(q[0]), q[1].strip()
                q_max_score = 10  # Default max score
            else:
                continue  # Skip invalid questions
            error_results.append({
                "question_number": q_num, 
                "status": "Error", 
                "score": 0, 
                "max_score": q_max_score, 
                "feedback": f"API error: {e}"
            })
        return error_results

def _generate_summary_with_gpt4o(full_analysis, rubric):
    # Use the session state summary prompt
    prompt = st.session_state.summary_prompt.format(
        full_analysis=json.dumps(full_analysis, indent=2),
        rubric=""
    )
    
    # Debug: Show the actual summary prompt being used (optional)
    if st.session_state.get('debug_mode', False):
        st.subheader("🔍 Debug: Summary Prompt Being Used")
        st.text_area("Summary Prompt", prompt, height=200, disabled=True)
    try:
        from openai import OpenAI
        client = OpenAI(api_key=OPENAI_API_KEY)
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        return json.loads(response.choices[0].message.content)
    except Exception:
        return {"overall_remarks": "Summary generation failed.", "rubric_analysis": []}

def grade_handwritten_submission_with_gpt4o(question_text, answer_images, rubric, test_title, progress_bar):
    # Get questions from the global question database
    questions = get_all_questions_for_test(test_title)
    
    if questions:
        st.info(f"✅ Using questions from question database ({len(questions)} questions)")
    else:
        # Try to parse from the question_text parameter
        if question_text and len(question_text.strip()) > 10:
            questions = re.findall(r'(\d+)[).]\s(.*?)(?:\[(\d+)\])', question_text, re.DOTALL)
            if questions:
                st.info(f"✅ Parsed {len(questions)} questions from question text")
                # Store the parsed questions in the database for future use
                for q_num, q_text, q_max_score in questions:
                    update_question_in_database(test_title, q_num, q_text, q_max_score)
            else:
                # Try alternative parsing patterns
                questions = re.findall(r'(\d+)[).]\s(.*?)(?=\d+[).]|$)', question_text, re.DOTALL)
                if questions:
                    st.info(f"✅ Parsed {len(questions)} questions using alternative pattern")
                    # Store with default max score of 10
                    for q_num, q_text in questions:
                        update_question_in_database(test_title, q_num, q_text.strip(), 10)
                else:
                    st.error("Could not parse questions from the question paper.")
                    st.info("💡 **Debug Info:** Question text length: " + str(len(question_text)) if question_text else "0")
                    st.info("💡 **Debug Info:** Question text preview: " + (question_text[:200] + "..." if question_text and len(question_text) > 200 else question_text or "None"))
                    return {"error": "Failed to parse questions."}
        else:
            st.error("No question text available for parsing.")
            st.info("💡 **Debug Info:** Question text is empty or too short")
            
            # Try to extract questions from answer images as a last resort
            st.info("🔄 Attempting to extract questions from answer images...")
            all_answer_text = ""
            for i, img in enumerate(answer_images):
                extracted_text = extract_text_from_image(img)
                if extracted_text and not extracted_text.startswith("GPT-4o OCR failed"):
                    all_answer_text += f"\n{extracted_text}\n"
            
            if all_answer_text:
                # Try to find questions in the answer text
                questions = re.findall(r'(\d+)[).]\s(.*?)(?=\d+[).]|$)', all_answer_text, re.DOTALL)
                if questions:
                    st.info(f"✅ Found {len(questions)} questions in answer images")
                    # Store with default max score of 10
                    for q_num, q_text in questions:
                        update_question_in_database(test_title, q_num, q_text.strip(), 10)
                else:
                    st.error("Could not find questions in answer images either.")
                    return {"error": "Failed to parse questions from any source."}
            else:
                st.error("Could not extract text from answer images.")
                return {"error": "No question text available."}
    
    # Extract full text from all answer images using custom OCR prompts
    extracted_answers = []
    for i, img in enumerate(answer_images):
        # Try to find the most relevant custom OCR prompt for this image
        # For now, use the default prompt, but this could be enhanced to match specific questions
        extracted_text = extract_text_from_image(img)
        extracted_answers.append({
            "image_number": i + 1,
            "extracted_text": extracted_text,
            "character_count": len(extracted_text) if extracted_text else 0
        })
    
    # Display parsing results
    st.subheader("🔍 Question Parsing Results:")
    st.info(f"✅ Successfully parsed {len(questions)} questions from question paper")
    
    with st.expander("📋 Parsed Questions", expanded=False):
        for i, q in enumerate(questions):
            # Handle tuples with different lengths (with or without max score)
            if len(q) == 3:
                q_num, q_text, q_max_score = int(q[0]), q[1].strip(), int(q[2])
            elif len(q) == 2:
                q_num, q_text = int(q[0]), q[1].strip()
                q_max_score = 10  # Default max score
            else:
                st.error(f"Invalid question format: {q}")
                continue
            st.write(f"**Question {q_num}** (Max Score: {q_max_score})")
            st.text_area(f"Question {q_num} Text", q_text, height=80, disabled=True)
            st.divider()
    
    # Show grading and summary prompts in debug mode (OCR prompt removed)
    if st.session_state.get('debug_mode', False):
        st.subheader("🔍 Debug: System Prompts")
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**📝 Grading Prompt:**")
            st.text_area("Grading", st.session_state.grading_prompt, height=200, disabled=True)
        
        with col2:
            st.markdown("**📊 Summary Prompt:**")
            st.text_area("Summary", st.session_state.summary_prompt, height=200, disabled=True)
    
    # Grade all questions at once
    progress_bar.progress(0.3, text="Preparing comprehensive analysis...")
    
    # Debug: Show which questions are being used for grading
    if st.session_state.get('debug_mode', False):
        st.subheader("🔍 Debug: Questions Being Used for Grading")
        for q_num, q_text, q_max_score in questions:
            st.write(f"**Q{q_num}** ({q_max_score}pts): {q_text}")
        
        st.subheader("🔍 Debug: Question Database Contents")
        for key, value in st.session_state.question_database.items():
            if key.startswith(f"{test_title.replace(' ', '_').replace('-', '_').replace('(', '').replace(')', '')}_Q"):
                st.write(f"**{key}**: {value['question_text']} (Modified: {value['last_modified']})")
    
    # Always show the questions being used for grading (not just in debug mode)
    st.subheader("📋 Questions Being Used for Evaluation:")
    for q in questions:
        # Handle tuples with different lengths (with or without max score)
        if len(q) == 3:
            q_num, q_text, q_max_score = int(q[0]), q[1].strip(), int(q[2])
        elif len(q) == 2:
            q_num, q_text = int(q[0]), q[1].strip()
            q_max_score = 10  # Default max score
        else:
            st.error(f"Invalid question format: {q}")
            continue
        st.write(f"**Q{q_num}** ({q_max_score}pts): {q_text}")
        st.divider()
    
    question_analysis = _grade_all_questions_with_gpt4o(questions, answer_images)
    
    # Calculate final scores
    total_raw_score = sum(q.get('score', 0) for q in question_analysis)
    total_possible_score = sum(q.get('max_score', 0) for q in question_analysis)
    final_percentage = math.ceil((total_raw_score / total_possible_score) * 100) if total_possible_score > 0 else 0
    final_analysis = {"total_score": final_percentage, "question_analysis": question_analysis}

    progress_bar.progress(0.8, text="Generating comprehensive analysis...")
    summary_data = _generate_summary_with_gpt4o(final_analysis, "")
    final_analysis.update(summary_data)
    
    # Save extracted content to file
    progress_bar.progress(0.9, text="Saving extracted content...")
    saved_file = save_extracted_content(
        test_title=test_title,
        student_name="Student",  # This will be updated when we have student info
        question_text=question_text,
        rubric="",
        answer_images=answer_images,
        question_analysis=question_analysis,
        final_score=final_percentage,
        summary_data=summary_data
    )
    
    if saved_file:
        final_analysis["saved_file"] = saved_file
    
    # Add extracted answers to the final analysis
    final_analysis["extracted_answers"] = extracted_answers
    
    progress_bar.progress(1.0, text="Grading completed!")
    return final_analysis

# --- UI Sections ---
with st.expander("➕ Add a New Test", expanded=True):
    with st.form("new_test_form", clear_on_submit=True):
        st.write("Define the test details, rubric, and upload the question paper.")
        test_title = st.text_input("Test Title", "Differentiation Test 2")
        test_date = st.date_input("Test Date")
        # Rubric input (PDF or image)
        rubric_upload = st.file_uploader("Upload Rubric (PDF or image)", type=["pdf", "png", "jpg", "jpeg"], key="rubric_uploader")
        question_paper_file = st.file_uploader("Upload Question Paper", type=["pdf", "png", "jpg", "jpeg"], key="question_uploader", help="Upload PDF or image files containing the question paper")
        submitted = st.form_submit_button("Create Test")
        if submitted:
            # Extract rubric text from upload, if provided
            test_rubric = ""
            rubric_images = []
            if rubric_upload is not None:
                mime = getattr(rubric_upload, 'type', '') or ''
                name = getattr(rubric_upload, 'name', '') or ''
                is_pdf = (mime == "application/pdf") or name.lower().endswith(".pdf")
                try:
                    if is_pdf:
                        # We need two passes: text and images for later OCR regeneration
                        test_rubric = extract_text_from_pdf(rubric_upload) or ""
                        try:
                            rubric_upload.seek(0)
                        except Exception:
                            pass
                        imgs = convert_pdf_to_images(rubric_upload)
                        rubric_images = imgs or []
                    elif mime.startswith("image/") or name.lower().endswith((".png", ".jpg", ".jpeg")):
                        image = Image.open(rubric_upload)
                        test_rubric = extract_text_from_image(image) or ""
                        rubric_images = [image]
                except Exception as e:
                    st.error(f"Rubric extraction failed: {e}")

            # Resolve question paper text
            question_text = ""
            if question_paper_file is not None:
                qp_mime = getattr(question_paper_file, 'type', '') or ''
                qp_name = getattr(question_paper_file, 'name', '') or ''
                if qp_mime == "application/pdf" or qp_name.lower().endswith('.pdf'):
                    question_text = extract_text_from_pdf(question_paper_file) or ""
                elif qp_mime.startswith("image/") or qp_name.lower().endswith((".png", ".jpg", ".jpeg")):
                    image = Image.open(question_paper_file)
                    question_text = extract_text_from_image(image) or ""
                if not question_text or question_text.startswith("GPT-4o OCR failed"):
                        question_text = "Question paper uploaded as image. Please ensure all questions are clearly visible."
            else:
                    st.error("Unsupported question paper file type. Please upload a PDF or image.")

            # Validate
            if not test_title:
                st.error("Please enter a test title.")
            if not test_date:
                st.error("Please select a test date.")
            # Rubric optional
            if not question_text:
                st.error("Question paper is required (PDF or image).")

            # Display extracted question text
            if question_text:
                # Parse and display individual questions with multiple patterns
                questions = []
                
                # Pattern 1: "1) Question text [3]" (with max score)
                pattern1 = re.findall(r'(\d+)[).]\s(.*?)(?:\[(\d+)\])', question_text, re.DOTALL)
                if pattern1:
                    questions = pattern1
                    st.info(f"✅ Parsed {len(questions)} questions using pattern 1 (with max scores)")
                
                # Pattern 2: "1) Question text" (without max score, default to 10)
                if not questions:
                    pattern2 = re.findall(r'(\d+)[).]\s(.*?)(?=\d+[).]|$)', question_text, re.DOTALL)
                    if pattern2:
                        questions = [(q[0], q[1].strip(), 10) for q in pattern2]  # Default max score of 10
                        st.info(f"✅ Parsed {len(questions)} questions using pattern 2 (default max score: 10)")
                
                # Pattern 3: "1. Question text" (with dot instead of parenthesis)
                if not questions:
                    pattern3 = re.findall(r'(\d+)\.\s(.*?)(?=\d+\.|$)', question_text, re.DOTALL)
                    if pattern3:
                        questions = [(q[0], q[1].strip(), 10) for q in pattern3]  # Default max score of 10
                        st.info(f"✅ Parsed {len(questions)} questions using pattern 3 (dot format)")
                
                # Pattern 4: "Question 1: Question text" (explicit question format)
                if not questions:
                    pattern4 = re.findall(r'Question\s+(\d+)[:.]\s*(.*?)(?=Question\s+\d+[:.]|$)', question_text, re.DOTALL | re.IGNORECASE)
                    if pattern4:
                        questions = [(q[0], q[1].strip(), 10) for q in pattern4]  # Default max score of 10
                        st.info(f"✅ Parsed {len(questions)} questions using pattern 4 (Question format)")
                
                # Debug: Show what we're trying to parse
                if not questions:
                    st.error("❌ Could not parse questions with any pattern")
                    st.write("**Debug: Question text to parse:**")
                    st.text_area("Raw Question Text", question_text, height=200, disabled=True)
                    st.write("**Tried patterns:**")
                    st.write("1. `1) Question text [3]` (with max score)")
                    st.write("2. `1) Question text` (without max score)")
                    st.write("3. `1. Question text` (dot format)")
                    st.write("4. `Question 1: Question text` (explicit format)")
                    
                    # Manual question creation option
                    st.subheader("🔧 Manual Question Creation")
                    st.info("Since automatic parsing failed, you can manually create questions:")
                    
                    num_questions = st.number_input("Number of questions to create:", min_value=1, max_value=20, value=3)
                    
                    manual_questions = []
                    for i in range(num_questions):
                        with st.container(border=True):
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                question_text_manual = st.text_area(f"Question {i+1} Text", height=100, key=f"manual_q{i+1}")
                            with col2:
                                max_score_manual = st.number_input(f"Max Score Q{i+1}", min_value=1, max_value=50, value=10, key=f"manual_score{i+1}")
                            
                            if question_text_manual.strip():
                                manual_questions.append((str(i+1), question_text_manual.strip(), max_score_manual))
                    
                    if manual_questions:
                        questions = manual_questions
                        st.success(f"✅ Created {len(questions)} manual questions")
                
                if questions:
                    # Removed detailed 'Individual Questions' preview section in Create Test
                    # Store questions in session state for persistent display and grading
                    # Use a sanitized key to avoid issues with special characters
                    sanitized_title = test_title.replace(" ", "_").replace("-", "_").replace("(", "").replace(")", "")
                    st.session_state[f"questions_{sanitized_title}"] = questions
                    
                    # Also store the original question text for reference
                    st.session_state[f"original_question_text_{sanitized_title}"] = question_text
                    
                    # Store all questions in the global question database
                    for q_num, q_text, q_max_score in questions:
                        update_question_in_database(test_title, q_num, q_text, q_max_score)
                    
                    for i, q in enumerate(questions):
                        q_num, q_text, q_max_score = int(q[0]), q[1].strip(), int(q[2])
                        
                        with st.container(border=True):
                            col1, col2 = st.columns([4, 1])
                            with col1:
                                st.markdown(f"**Question {q_num}** (Max Score: {q_max_score})")
                                st.text_area(f"Question {q_num} Text", q_text, height=100, disabled=True, key=f"question_display_create_{test_title}_{q_num}")
                            with col2:
                                # Individual OCR prompt for each question
                                st.markdown("**OCR Prompt:**")
                                default_ocr_prompt = f"TRANSCRIBE QUESTION {q_num} FROM THIS IMAGE. Output only the question text you see."
                                custom_ocr_prompt = st.text_area(
                                    f"OCR Prompt for Q{q_num}",
                                    value=default_ocr_prompt,
                                    height=120,
                                    key=f"ocr_prompt_create_{test_title}_{q_num}",
                                    help=f"Custom OCR prompt for Question {q_num}"
                                )
                                # Store the custom OCR prompt
                                st.session_state[f"custom_ocr_prompt_{sanitized_title}_{q_num}"] = custom_ocr_prompt
            else:
                st.error("❌ Failed to extract text from question paper")
            
            # Removed inline rubric preview section; rubric will be shown in the test's question view
            
            if question_text:
                new_test = {"id": f"test_{len(st.session_state.tests) + 1}", "title": test_title, "date": str(test_date), "rubric": test_rubric or "", "question_text": question_text}
                st.session_state.tests.append(new_test)
                # Persist rubric images and parsed rows for later per-concept OCR/edits
                try:
                    parsed_rows = extract_rubric_table_from_text(test_rubric) if (test_rubric and len(test_rubric) > 0) else []
                    st.session_state[f"rubric_rows_{new_test['id']}"] = parsed_rows
                except Exception:
                    st.session_state[f"rubric_rows_{new_test['id']}"] = []
                if rubric_images:
                    st.session_state[f"rubric_images_{new_test['id']}"] = rubric_images
                st.success(f"Test '{test_title}' created!")
        elif submitted:
            st.warning("Please fill out all fields and upload required files.")

st.divider()

# Display persistent questions for existing tests
if st.session_state.tests:
    st.subheader("📋 Test Questions & OCR Configuration")
    st.info(f"Found {len(st.session_state.tests)} test(s) in session state")
    
    for test_idx, test in enumerate(st.session_state.tests):
        # Use sanitized title to match the stored questions
        sanitized_title = test['title'].replace(" ", "_").replace("-", "_").replace("(", "").replace(")", "")
        questions = st.session_state.get(f"questions_{sanitized_title}", [])
        
        st.write(f"Debug: Test '{test['title']}' has {len(questions) if questions else 0} questions")
        
        if questions:
            with st.expander(f"📝 {test['title']} - Questions & OCR Settings", expanded=True):
                st.info(f"Configure individual OCR prompts for each question in {test['title']}")
                
                # Rubric display as table (if we have extracted rubric text)
                if test.get('rubric'):
                    # Use cached parsed rows if present, else parse
                    cached_rows_key = f"rubric_rows_{test['id']}"
                    rubric_rows = st.session_state.get(cached_rows_key)
                    if not rubric_rows:
                        # Text-only parse first (robust)
                        rubric_rows_text = parse_rubric_text_llm(test.get('rubric', ''))
                        rubric_rows_text = fill_missing_examples_from_text(rubric_rows_text, test.get('rubric', ''))

                        # Also attempt image-based parse to capture formulas embedded as images
                        rub_imgs_key = f"rubric_images_{test['id']}"
                        available_imgs = st.session_state.get(rub_imgs_key, [])
                        rubric_rows_img = extract_rubric_rows_from_images(available_imgs) if available_imgs else []
                        # Merge: prefer image Examples when available
                        rubric_rows = merge_rubric_rows_by_concept(rubric_rows_text, rubric_rows_img)
                    if not rubric_rows:
                        rubric_rows = extract_rubric_table_from_text(test['rubric'])
                    st.session_state[cached_rows_key] = rubric_rows
                    if rubric_rows:
                        # Debug: show scanning status for Examples
                        empty_examples = sum(1 for r in rubric_rows if not str(r.get('Example', '')).strip())
                        total_rows = len(rubric_rows)
                        st.caption(f"Rubric rows loaded: {total_rows}. Empty Examples: {empty_examples}.")
                        st.markdown("**Rubric (editable)**")
                        # In-place editable table for the rubric
                        import pandas as pd
                        editor_key = f"rubric_editor_{test['id']}"
                        df = pd.DataFrame(rubric_rows)
                        edited_df = st.data_editor(
                            df,
                            use_container_width=True,
                            num_rows="dynamic",
                            column_config={
                                "Concept No.": st.column_config.NumberColumn("Concept No.", step=1, help="Sequential concept number"),
                                "Concept": st.column_config.TextColumn("Concept", help="Short concept title"),
                                "Example": st.column_config.TextColumn("Example", help="Brief example or formula"),
                                "Status": st.column_config.TextColumn("Status", help="Optional status/notes"),
                            },
                            key=editor_key,
                        )

                        # Debug visibility: show extracted rubric data (no nested expander)
                        cols_dbg = st.columns(3)
                        if cols_dbg[0].button("Show parsed JSON", key=f"show_rubric_json_{test['id']}"):
                            st.json(rubric_rows)
                        if cols_dbg[1].button("Show raw rubric text", key=f"show_rubric_text_{test['id']}"):
                            st.text_area("Raw Rubric Text", test.get('rubric',''), height=240, disabled=True, key=f"raw_rubric_text_{test['id']}")
                        if cols_dbg[2].button("Copy JSON", key=f"copy_rubric_json_{test['id']}"):
                            st.code(json.dumps(rubric_rows, ensure_ascii=False, indent=2))

                        cols = st.columns(3)
                        if cols[0].button("💾 Save Rubric", key=f"save_rubric_table_{test['id']}"):
                            # Persist edited rows back to session with clean numbering/types
                            rows = edited_df.to_dict(orient='records') if hasattr(edited_df, 'to_dict') else edited_df
                            cleaned = []
                            for idx, r in enumerate(rows, start=1):
                                cleaned.append({
                                    "Concept No.": int(r.get("Concept No.") or idx),
                                    "Concept": str(r.get("Concept", "")).strip(),
                                    "Example": str(r.get("Example", "")).strip(),
                                    "Status": str(r.get("Status", "")).strip(),
                                })
                            st.session_state[cached_rows_key] = cleaned
                            st.success("Rubric saved.")

                        if cols[1].button("↩️ Reset", key=f"reset_rubric_table_{test['id']}"):
                                    fresh = extract_rubric_table_from_text(test['rubric'])
                                    st.session_state[cached_rows_key] = fresh
                                    st.success("Reset to parsed rubric.")

                        if cols[2].button("➕ Add 5 Rows", key=f"add_rows_rubric_{test['id']}"):
                            # Extend with blank rows for convenience
                            rows = edited_df.to_dict(orient='records') if hasattr(edited_df, 'to_dict') else rubric_rows
                            for _ in range(5):
                                rows.append({"Concept No.": None, "Concept": "", "Example": "", "Status": ""})
                            st.session_state[cached_rows_key] = rows
                            st.rerun()
                

                
                for i, q in enumerate(questions):
                    q_num, q_text, q_max_score = int(q[0]), q[1].strip(), int(q[2])
                    
                    with st.container(border=True):
                        col1, col2 = st.columns([4, 1])
                        with col1:
                            st.markdown(f"**Question {q_num}** (Max Score: {q_max_score})")
                            
                            # Get the current question text from database (or use original if not found)
                            current_question_text = get_question_from_database(test['title'], q_num) or q_text
                            
                            # Check if we're editing this question
                            if st.button("✏️ Edit Question", key=f"edit_persistent_question_{test_idx}_{sanitized_title}_{q_num}"):
                                st.session_state[f"editing_persistent_question_{test_idx}_{sanitized_title}_{q_num}"] = True
                            
                            # Show question text (editable or read-only)
                            if st.session_state.get(f"editing_persistent_question_{test_idx}_{sanitized_title}_{q_num}", False):
                                edited_question = st.text_area(
                                    f"Question {q_num} Text (Editing)",
                                    value=current_question_text,
                                    height=100,
                                    key=f"editing_persistent_question_text_{test_idx}_{sanitized_title}_{q_num}"
                                )
                                
                                col1, col2 = st.columns(2)
                                if col1.button("✅ Save Question", key=f"save_persistent_question_{test_idx}_{sanitized_title}_{q_num}"):
                                    # Update the question in database
                                    update_question_in_database(test['title'], q_num, edited_question, q_max_score)
                                    st.success(f"Question {q_num} updated!")
                                    st.session_state[f"editing_persistent_question_{test_idx}_{sanitized_title}_{q_num}"] = False
                                    st.rerun()
                                if col2.button("❌ Cancel", key=f"cancel_persistent_question_{test_idx}_{sanitized_title}_{q_num}"):
                                    st.session_state[f"editing_persistent_question_{test_idx}_{sanitized_title}_{q_num}"] = False
                                    st.rerun()
                            else:
                                st.text_area(f"Question {q_num} Text", current_question_text, height=100, disabled=True, key=f"persistent_question_{test_idx}_{sanitized_title}_{q_num}")
                        with col2:
                            # Individual OCR prompt for each question
                            st.markdown("**OCR Prompt:**")
                            default_ocr_prompt = (
                "You are an elite OCR agent for handwritten STEM exams. "
                f"Extract EXACTLY the problem statement for Question {q_num} from the provided image(s). "
                "Be robust to noise, skew, shadows, and writing styles.\n\n"
                "Rules:\n"
                "- Output: plain text only. No preface, no extra lines.\n"
                "- Scope: include only the question text and any labeled subparts ((a), (b), ...). Ignore unrelated headers/footers.\n"
                "- Math fidelity: preserve all math. Use readable inline forms: ∫, Σ, Π, d/dx, lim, |x|, sqrt(), ^ for exponent, / for fractions, () for grouping.\n"
                "- Structure: keep original line breaks and indentation for subparts.\n"
                "- Normalization: unify symbols when ambiguous: '*' for multiplication, '^' for powers, '/' for rational forms.\n"
                "- Error handling: never invent content. If unreadable, write '[illegible]'. If a symbol is ambiguous, choose the most likely and add the alternative in brackets, e.g., 'x^2 [or x^z]'.\n"
                "- Cleanup: remove page numbers, watermarks, and scribbles. Expand common abbreviations when unambiguous (e.g., 'w.r.t.' -> 'with respect to').\n"
                "- Language: keep the original language.\n\n"
                f"Context hint (first 120 chars of parsed text): {q_text[:120]}"
                            )
                            stored_prompt = st.session_state.get(
                                f"custom_ocr_prompt_{sanitized_title}_{q_num}",
                                default_ocr_prompt,
                            )
                            custom_ocr_prompt = st.text_area(
                                f"OCR Prompt for Q{q_num}",
                                value=stored_prompt,
                                height=120,
                                key=f"persistent_ocr_prompt_{test_idx}_{sanitized_title}_{q_num}",
                                help=f"Custom OCR prompt for extracting Question {q_num} from question paper images",
                            )
                            # Store the custom OCR prompt
                            st.session_state[
                                f"custom_ocr_prompt_{sanitized_title}_{q_num}"
                            ] = custom_ocr_prompt
                            
                            # Simple regenerate button that modifies the existing text
                            if st.button("🔄 Regenerate Question", key=f"regenerate_question_{test_idx}_{sanitized_title}_{q_num}"):
                                # Set flag to show processing
                                st.session_state[f"processing_question_{test_idx}_{sanitized_title}_{q_num}"] = True
                                st.rerun()
                            
                            # Show processing and results
                            if st.session_state.get(f"processing_question_{test_idx}_{sanitized_title}_{q_num}", False):
                                with st.spinner(f"Modifying Question {q_num} with custom OCR prompt..."):
                                    try:
                                        from openai import OpenAI
                                        client = OpenAI(api_key=OPENAI_API_KEY)
                                        
                                                                                # Create a prompt that applies the custom OCR instructions to the existing text
                                        processing_prompt = f"""
You are an expert at converting complex mathematical notation to human-readable mathematical notation.

OCR Instructions: {custom_ocr_prompt}

Original Question Text: {q_text}

Please convert to human-readable mathematical notation:
- Keep mathematical symbols but make them clear and readable
- Convert sqrt(a) to √a (square root symbol)
- Convert sin^-1(t) to sin⁻¹(t) or arcsin(t)
- Convert cos^-1(t) to cos⁻¹(t) or arccos(t)
- Convert dy/dx to dy/dx (keep as fraction)
- Convert fractions like a/b to a/b (keep as fraction)
- Use proper mathematical notation that's easy to read
- Keep it mathematical but readable
- No backslashes or dollar signs

Return only the human-readable mathematical question text, no explanations or additional text.
"""
                                        
                                        response = client.chat.completions.create(
                                            model="gpt-4o",
                                            messages=[
                                                {"role": "user", "content": processing_prompt}
                                            ],
                                            temperature=0.0
                                        )
                                        
                                        modified_question = response.choices[0].message.content.strip()
                                        
                                        # Store the results for display
                                        st.session_state[f"original_question_{test_idx}_{sanitized_title}_{q_num}"] = q_text
                                        st.session_state[f"modified_question_{test_idx}_{sanitized_title}_{q_num}"] = modified_question
                                        
                                        # Update the question in the global database
                                        update_question_in_database(test['title'], q_num, modified_question, q_max_score)
                                        
                                        # Auto-regenerate analysis for all students if question was modified
                                        if st.session_state.students and st.session_state.submissions:
                                            st.info("🔄 Auto-regenerating analysis for all students with the updated question...")
                                            
                                            for student in st.session_state.students:
                                                submission_key = f"test_{test['id']}_{student['id']}"
                                                submission = st.session_state.submissions.get(submission_key)
                                                
                                                if submission and submission.get('extracted_answers'):
                                                    # Get the updated question text
                                                    updated_question_text = modified_question
                                                    
                                                    # Extract student answer for this question
                                                    student_answer = extract_question_answer(submission['extracted_answers'], q_num)
                                                    
                                                    if student_answer and student_answer.strip():
                                                        # Regenerate analysis with the updated question
                                                        new_analysis, message = regenerate_analysis_for_question(
                                                            q_num, "", updated_question_text, student_answer, test['rubric'], test['id'], student['id'], submission['extracted_answers']
                                                        )
                                                        
                                                        if new_analysis:
                                                            # Update the question analysis in the submission
                                                            for i, qa_item in enumerate(submission.get('question_analysis', [])):
                                                                if qa_item.get('question_number') == q_num:
                                                                    submission['question_analysis'][i].update(new_analysis)
                                                                    break
                                                            
                                                            # Recalculate total score
                                                            total_raw_score = sum(q.get('score', 0) for q in submission.get('question_analysis', []))
                                                            total_possible_score = sum(q.get('max_score', 0) for q in submission.get('question_analysis', []))
                                                            final_percentage = math.ceil((total_raw_score / total_possible_score) * 100) if total_possible_score > 0 else 0
                                                            submission['total_score'] = final_percentage
                                                            
                                                            # Save to database
                                                            save_updated_results_to_database(test['id'], student['id'], q_num, new_analysis)
                                        
                                        # Clear processing flag
                                        st.session_state[f"processing_question_{test_idx}_{sanitized_title}_{q_num}"] = False
                                        st.success(f"Question {q_num} updated and analysis regenerated for all students!")
                                        st.rerun()
                                        
                                    except Exception as e:
                                        st.error(f"Error applying OCR prompt: {str(e)}")
                                        st.session_state[f"processing_question_{test_idx}_{sanitized_title}_{q_num}"] = False
                                        st.rerun()
                            
                            # Show before/after comparison if available
                            if st.session_state.get(f"modified_question_{test_idx}_{sanitized_title}_{q_num}"):
                                st.markdown("**🔄 Question Modification Results:**")
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.markdown("**Before:**")
                                    original = st.session_state.get(f"original_question_{test_idx}_{sanitized_title}_{q_num}", q_text)
                                    st.text_area(f"Original Q{q_num}", original, height=100, disabled=True)
                                with col2:
                                    st.markdown("**After:**")
                                    modified = st.session_state.get(f"modified_question_{test_idx}_{sanitized_title}_{q_num}", q_text)
                                    st.text_area(f"Modified Q{q_num}", modified, height=100, disabled=True)
                                
                                col1, col2 = st.columns(2)
                                if col1.button("✅ Accept Changes", key=f"accept_changes_{test_idx}_{sanitized_title}_{q_num}"):
                                    # Clear the comparison display
                                    del st.session_state[f"original_question_{test_idx}_{sanitized_title}_{q_num}"]
                                    del st.session_state[f"modified_question_{test_idx}_{sanitized_title}_{q_num}"]
                                    st.rerun()
                                if col2.button("🔄 Regenerate Analysis", key=f"regenerate_after_change_{test_idx}_{sanitized_title}_{q_num}"):
                                    # Manually regenerate analysis for all students with the updated question
                                    if st.session_state.students and st.session_state.submissions:
                                        with st.spinner(f"🔄 Regenerating analysis for Question {q_num}..."):
                                            updated_question_text = get_question_from_database(test['title'], q_num) or q_text
                                            
                                            for student in st.session_state.students:
                                                submission_key = f"test_{test['id']}_{student['id']}"
                                                submission = st.session_state.submissions.get(submission_key)
                                                
                                                if submission and submission.get('extracted_answers'):
                                                    student_answer = extract_question_answer(submission['extracted_answers'], q_num)
                                                    
                                                    if student_answer and student_answer.strip():
                                                        new_analysis, message = regenerate_analysis_for_question(
                                                            q_num, "", updated_question_text, student_answer, test['rubric'], test['id'], student['id'], submission['extracted_answers']
                                                        )
                                                        
                                                        if new_analysis:
                                                            for i, qa_item in enumerate(submission.get('question_analysis', [])):
                                                                if qa_item.get('question_number') == q_num:
                                                                    submission['question_analysis'][i].update(new_analysis)
                                                                    break
                                                            
                                                            # Recalculate total score
                                                            total_raw_score = sum(q.get('score', 0) for q in submission.get('question_analysis', []))
                                                            total_possible_score = sum(q.get('max_score', 0) for q in submission.get('question_analysis', []))
                                                            final_percentage = math.ceil((total_raw_score / total_possible_score) * 100) if total_possible_score > 0 else 0
                                                            submission['total_score'] = final_percentage
                                                            
                                                            save_updated_results_to_database(test['id'], student['id'], q_num, new_analysis)
                                            
                                            st.success(f"Analysis regenerated for Question {q_num}!")
                                            st.rerun()
        else:
            st.warning(f"⚠️ No questions found for test '{test['title']}'. Questions may not have been parsed correctly during test creation.")

st.divider()
if not st.session_state.tests:
    st.info("No tests created yet. Add a new test to get started.")
else:
    for test in reversed(st.session_state.tests):
        with st.container(border=True):
            st.header(f"📝 Test: {test['title']}")
            st.caption(f"Date: {test['date']}")
            for student in st.session_state.students:
                submission_key = f"test_{test['id']}_{student['id']}"
                submission = st.session_state.submissions.get(submission_key)
                cols = st.columns([2, 1, 3])
                cols[0].write(f"**Student:** {student['name']}")
                if submission:
                    score = submission.get("total_score", "N/A")
                    if 'error' in submission:
                        cols[1].error("Error")
                    else:
                        cols[1].success(f"Graded: {score}/100")
                    with cols[2].popover("View Full Analysis"):
                        st.subheader("Detailed Analysis")
                        
                        # --- UI FIX IS HERE ---
                        st.markdown("**Overall Remarks**")
                        remarks = submission.get('overall_remarks', 'No summary provided.')
                        st.markdown(f"> {remarks}") # Using markdown for better text flow
                        # --- END OF UI FIX ---

                        # Display Questions and Student Answers
                        st.markdown("**📋 Questions and Student Answers**")
                        
                        # Get questions from the database first, then fall back to parsing
                        sanitized_title = test['title'].replace(" ", "_").replace("-", "_").replace("(", "").replace(")", "")
                        questions = get_all_questions_for_test(test['title'])
                        
                        if not questions:
                            # Fall back to parsing from question text
                            if test['question_text'] and len(test['question_text'].strip()) > 10:
                                questions = re.findall(r'(\d+)[).]\s(.*?)(?:\[(\d+)\])', test['question_text'], re.DOTALL)
                                if not questions:
                                    # Try alternative parsing patterns
                                    questions = re.findall(r'(\d+)[).]\s(.*?)(?=\d+[).]|$)', test['question_text'], re.DOTALL)
                                    if questions:
                                        # Convert to proper format with default max score
                                        questions = [(int(q[0]), q[1].strip(), 10) for q in questions]
                        
                        if questions:
                            for i, q in enumerate(questions):
                                # Handle tuples with different lengths (with or without max score)
                                if len(q) == 3:
                                    q_num, q_text, q_max_score = int(q[0]), q[1].strip(), int(q[2])
                                elif len(q) == 2:
                                    q_num, q_text = int(q[0]), q[1].strip()
                                    q_max_score = 10  # Default max score
                                else:
                                    st.error(f"Invalid question format: {q}")
                                    continue
                                
                                with st.container(border=True):
                                    st.markdown(f"**Question {q_num}** (Max Score: {q_max_score})")
                                    
                                    # Question text with edit capability
                                    col1, col2 = st.columns([4, 1])
                                    with col1:
                                        st.text_area(f"Question {q_num} Text", q_text, height=80, disabled=True, key=f"question_text_{test['id']}_{student['id']}_{q_num}")
                                    with col2:
                                        if st.button("✏️ Edit Question", key=f"edit_question_{test['id']}_{student['id']}_{q_num}"):
                                            st.session_state[f"editing_question_{test['id']}_{student['id']}_{q_num}"] = True
                                    
                                    # Edit question text
                                    if st.session_state.get(f"editing_question_{test['id']}_{student['id']}_{q_num}", False):
                                        edited_question = st.text_area("Edit Question Text", q_text, height=100, key=f"edit_question_text_{test['id']}_{student['id']}_{q_num}")
                                        col1, col2 = st.columns(2)
                                        if col1.button("✅ Save", key=f"save_question_{test['id']}_{student['id']}_{q_num}"):
                                            st.success("Question updated!")
                                            st.session_state[f"editing_question_{test['id']}_{student['id']}_{q_num}"] = False
                                            st.rerun()
                                        if col2.button("❌ Cancel", key=f"cancel_question_{test['id']}_{student['id']}_{q_num}"):
                                            st.session_state[f"editing_question_{test['id']}_{student['id']}_{q_num}"] = False
                                            st.rerun()
                                    
                                    # Find corresponding question analysis
                                    qa = next((qa for qa in submission.get('question_analysis', []) if qa.get('question_number') == q_num), None)
                                    
                                    if qa:
                                        # Show complete extracted answer for this question
                                        extracted_answers = submission.get('extracted_answers', [])
                                        if extracted_answers:
                                            st.markdown("**Complete Student Answer (Extracted from Images):**")
                                            
                                            # Intelligent question answer extraction (respects overrides)
                                            relevant_text = get_current_answer_for_question(submission, q_num)
                                            
                                            if relevant_text.strip():
                                                col1, col2 = st.columns([4, 1])
                                                with col1:
                                                    st.text_area(f"Complete Answer for Q{q_num}", relevant_text.strip(), height=300, disabled=True, key=f"complete_answer_{test['id']}_{student['id']}_{q_num}")
                                                with col2:
                                                    if st.button("✏️ Edit Answer", key=f"edit_answer_{test['id']}_{student['id']}_{q_num}"):
                                                        st.session_state[f"editing_answer_{test['id']}_{student['id']}_{q_num}"] = True
                                                
                                                # Regenerate OCR for this question (answers) immediately using stored/default prompt
                                                if st.button("🔄 Regenerate Answer OCR", key=f"regenerate_answer_ocr_{test['id']}_{student['id']}_{q_num}"):
                                                    # Resolve prompt to use (persisted per test/student/question)
                                                    ocr_prompt_key = f"ocr_prompt_answer_{test['id']}_{student['id']}_{q_num}"
                                                    default_answer_ocr_prompt = (
                                                        "You are an elite OCR agent for handwritten exam answers. "
                                                        f"Extract the COMPLETE student answer for Question {q_num} from the provided image(s).\n\n"
                                                        "BOUNDARIES\n"
                                                        f"- START: the first line that begins with the answer number for Q{q_num} (e.g., '{q_num}.', '{q_num})', 'Q{q_num}', '(" + str(q_num) + ")').\n"
                                                        "- END: the first line that begins with the next answer number (Q{n+1}) — do NOT include that next number. If not found, end at the final line of the image set.\n\n"
                                                        "RULES\n"
                                                        "- Preserve ALL content and original line breaks. Capture steps before and after equations; do not summarize.\n"
                                                        "- Preserve math faithfully using readable inline forms: sqrt(), abs(), d/dx, ∫, Σ, Π, lim, |x|, ^ for powers, / for fractions, parentheses for grouping.\n"
                                                        "- Convert LaTeX tokens to readable math (e.g., \\frac{dy}{dx} -> d/dx, \\sqrt{x} -> sqrt(x)). Remove LaTeX backslashes and $.\n"
                                                        "- Ignore clearly crossed-out work only if fully struck; minor strike-offs should not erase valid surrounding content.\n"
                                                        "- If unreadable, use [illegible]. Do not invent content.\n"
                                                        "- Output plain text only; no headers, no labels, no JSON."
                                                    )
                                                    effective_prompt = st.session_state.get(ocr_prompt_key, default_answer_ocr_prompt)
                                                    with st.spinner("🔄 Regenerating OCR for answers..."):
                                                        new_extractions, message = regenerate_ocr_for_question(
                                                            q_num, effective_prompt, extracted_answers, test['id'], student['id']
                                                        )
                                                        if new_extractions:
                                                            submission['extracted_answers'] = new_extractions
                                                            st.success(f"Answer OCR regenerated. {message}")
                                                        else:
                                                            st.error(f"Failed to regenerate OCR: {message}")
                                                    st.rerun()

                                                # Advanced: open UI to customize OCR prompt and run
                                                if st.button("⚙️ Custom Answer OCR", key=f"open_custom_answer_ocr_{test['id']}_{student['id']}_{q_num}"):
                                                    st.session_state[f"regenerating_ocr_{test['id']}_{student['id']}_{q_num}"] = True
                                                
                                                # OCR regeneration interface (customization panel)
                                                if st.session_state.get(f"regenerating_ocr_{test['id']}_{student['id']}_{q_num}", False):
                                                    st.markdown("**🔄 Regenerate OCR for Question {q_num}**")
                                                    
                                                    # Custom OCR prompt for this question's answers
                                                    # Persist per test/student/question under session_state
                                                    ocr_prompt_key = f"ocr_prompt_answer_{test['id']}_{student['id']}_{q_num}"
                                                    default_answer_ocr_prompt = f"Extract the student's full answer for Question {q_num} from this image. Preserve math structure, convert LaTeX to readable math, remove crossed-out work, and keep steps order."
                                                    current_prompt_val = st.session_state.get(ocr_prompt_key, default_answer_ocr_prompt)
                                                    # IMPORTANT: Do not set session_state here; just control the widget's initial value
                                                    if ocr_prompt_key in st.session_state:
                                                        custom_ocr_prompt = st.text_area(
                                                            f"Custom Answer OCR Prompt for Q{q_num}",
                                                            height=120,
                                                            key=ocr_prompt_key
                                                        )
                                                    else:
                                                        custom_ocr_prompt = st.text_area(
                                                            f"Custom Answer OCR Prompt for Q{q_num}",
                                                            value=default_answer_ocr_prompt,
                                                            height=120,
                                                            key=ocr_prompt_key
                                                    )
                                                    
                                                    # Show original images for this question
                                                    st.markdown("**Original Images:**")
                                                    for answer in extracted_answers:
                                                        if answer.get('extracted_text'):
                                                            # Check if this image contains the question
                                                            if f"{q_num}." in answer['extracted_text'] or f"{q_num})" in answer['extracted_text']:
                                                                st.markdown(f"**Image {answer['image_number']}:**")
                                                                # Here you would show the actual image
                                                                st.text_area(f"Current OCR for Image {answer['image_number']}", answer['extracted_text'], height=150, disabled=True)
                                                    
                                                    col1, col2, col3 = st.columns(3)
                                                    if col1.button("🔄 Run Custom OCR", key=f"execute_custom_answer_ocr_{test['id']}_{student['id']}_{q_num}"):
                                                        with st.spinner("🔄 Regenerating OCR with custom prompt..."):
                                                            # Get the custom prompt
                                                            custom_prompt = st.session_state.get(ocr_prompt_key, current_prompt_val)
                                                            
                                                            # Debug information
                                                            if st.session_state.get('debug_mode', False):
                                                                st.write(f"Debug: Custom prompt: {custom_prompt}")
                                                                st.write(f"Debug: Question number: {q_num}")
                                                                st.write(f"Debug: Test ID: {test['id']}, Student ID: {student['id']}")
                                                                session_key = f'original_images_{test["id"]}_{student["id"]}'
                                                                st.write(f"Debug: Original images in session: {session_key in st.session_state}")
                                                                if session_key in st.session_state:
                                                                    st.write(f"Debug: Number of original images: {len(st.session_state[session_key])}")
                                                            
                                                            # Call the regeneration function
                                                            new_extractions, message = regenerate_ocr_for_question(
                                                                q_num, custom_prompt, extracted_answers, test['id'], student['id']
                                                            )
                                                            
                                                            if new_extractions:
                                                                # Merge replacements by image number
                                                                if isinstance(submission.get('extracted_answers'), list):
                                                                    existing = {a['image_number']: a for a in submission['extracted_answers']}
                                                                    for ne in new_extractions:
                                                                        existing[ne['image_number']] = ne
                                                                    submission['extracted_answers'] = list(existing.values())
                                                                else:
                                                                    submission['extracted_answers'] = new_extractions
                                                                    st.success(f"OCR regenerated successfully! {message}")
                                                                
                                                                # Debug: Show new extractions
                                                                if st.session_state.get('debug_mode', False):
                                                                    st.write("Debug: New extractions:")
                                                                    for ext in new_extractions:
                                                                        st.write(f"  Image {ext['image_number']}: {ext['extracted_text'][:100]}...")
                                                            else:
                                                                st.error(f"Failed to regenerate OCR: {message}")
                                                            
                                                            st.session_state[f"regenerating_ocr_{test['id']}_{student['id']}_{q_num}"] = False
                                                            st.rerun()
                                                    if col2.button("✅ Use Default Prompt", key=f"default_answer_ocr_run_{test['id']}_{student['id']}_{q_num}"):
                                                        with st.spinner("🔄 Regenerating OCR with default prompt..."):
                                                            new_extractions, message = regenerate_ocr_for_question(
                                                                q_num, default_answer_ocr_prompt, extracted_answers, test['id'], student['id']
                                                            )
                                                            
                                                            if new_extractions:
                                                                # Merge replacements by image number
                                                                if isinstance(submission.get('extracted_answers'), list):
                                                                    existing = {a['image_number']: a for a in submission['extracted_answers']}
                                                                    for ne in new_extractions:
                                                                        existing[ne['image_number']] = ne
                                                                    submission['extracted_answers'] = list(existing.values())
                                                                else:
                                                                    submission['extracted_answers'] = new_extractions
                                                                    st.success(f"OCR regenerated with default prompt! {message}")
                                                            else:
                                                                st.error(f"Failed to regenerate OCR: {message}")
                                                            
                                                            st.session_state[f"regenerating_ocr_{test['id']}_{student['id']}_{q_num}"] = False
                                                            st.rerun()
                                                    if col3.button("❌ Cancel OCR", key=f"cancel_ocr_{test['id']}_{student['id']}_{q_num}"):
                                                        st.session_state[f"regenerating_ocr_{test['id']}_{student['id']}_{q_num}"] = False
                                                        st.rerun()
                                                
                                                # Edit answer text
                                                if st.session_state.get(f"editing_answer_{test['id']}_{student['id']}_{q_num}", False):
                                                    edited_answer = st.text_area("Edit Student Answer", relevant_text.strip(), height=300, key=f"edit_answer_text_{test['id']}_{student['id']}_{q_num}")
                                                    col1, col2 = st.columns(2)
                                                    if col1.button("✅ Save Answer", key=f"save_answer_{test['id']}_{student['id']}_{q_num}"):
                                                        # Persist manual override on the submission object
                                                        submission_key = f"test_{test['id']}_{student['id']}"
                                                        if 'answer_overrides' not in submission:
                                                            submission['answer_overrides'] = {}
                                                        submission['answer_overrides'][str(q_num)] = edited_answer.strip()
                                                        # Mirror to global session submissions store if present
                                                        if submission_key in st.session_state.submissions:
                                                            st.session_state.submissions[submission_key].setdefault('answer_overrides', {})
                                                            st.session_state.submissions[submission_key]['answer_overrides'][str(q_num)] = edited_answer.strip()
                                                        st.success("Answer updated!")
                                                        st.session_state[f"editing_answer_{test['id']}_{student['id']}_{q_num}"] = False
                                                        st.rerun()
                                                    if col2.button("❌ Cancel Answer", key=f"cancel_answer_{test['id']}_{student['id']}_{q_num}"):
                                                        st.session_state[f"editing_answer_{test['id']}_{student['id']}_{q_num}"] = False
                                                        st.rerun()
                                            else:
                                                st.warning(f"No specific answer found for Question {q_num}")
                                        else:
                                            st.warning("No extracted answers available.")
                                        
                                        # Show AI's work summary with edit capability
                                        if qa.get('extracted_work'):
                                            st.markdown("**AI's Work Summary:**")
                                            col1, col2 = st.columns([4, 1])
                                            with col1:
                                                st.info(qa.get('extracted_work'))
                                            with col2:
                                                if st.button("✏️ Edit Summary", key=f"edit_summary_{test['id']}_{student['id']}_{q_num}"):
                                                    st.session_state[f"editing_summary_{test['id']}_{student['id']}_{q_num}"] = True
                                            
                                            # Edit work summary
                                            if st.session_state.get(f"editing_summary_{test['id']}_{student['id']}_{q_num}", False):
                                                edited_summary = st.text_area("Edit Work Summary", qa.get('extracted_work', ''), height=100, key=f"edit_summary_text_{test['id']}_{student['id']}_{q_num}")
                                                col1, col2 = st.columns(2)
                                                if col1.button("✅ Save Summary", key=f"save_summary_{test['id']}_{student['id']}_{q_num}"):
                                                    st.success("Work summary updated!")
                                                    st.session_state[f"editing_summary_{test['id']}_{student['id']}_{q_num}"] = False
                                                    st.rerun()
                                                if col2.button("❌ Cancel Summary", key=f"cancel_summary_{test['id']}_{student['id']}_{q_num}"):
                                                    st.session_state[f"editing_summary_{test['id']}_{student['id']}_{q_num}"] = False
                                                    st.rerun()
                                        
                                        # Show score and status with edit capability
                                        col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
                                        with col1:
                                            score = qa.get('score', 0)
                                            max_score = qa.get('max_score', 0)
                                            st.metric("Score", f"{score}/{max_score}")
                                        with col2:
                                            if st.button("✏️ Edit Score", key=f"edit_score_{test['id']}_{student['id']}_{q_num}"):
                                                st.session_state[f"editing_score_{test['id']}_{student['id']}_{q_num}"] = True
                                        with col3:
                                            status = qa.get('status', 'Unknown')
                                            status_color = {
                                                'Excellent': 'success',
                                                'Good': 'info', 
                                                'Fair': 'warning',
                                                'Poor': 'error',
                                                'Not Attempted': 'error'
                                            }.get(status, 'secondary')
                                            st.write(f"**Status:** :{status_color}[{status}]")
                                        with col4:
                                            # Check if this result was recently updated
                                            result_key = f"{test['id']}_{student['id']}_Q{q_num}"
                                            if "results_database" in st.session_state and result_key in st.session_state.results_database:
                                                updated_time = st.session_state.results_database[result_key]["updated_at"]
                                                st.info("🔄 Updated")
                                        
                                        # Edit score
                                        if st.session_state.get(f"editing_score_{test['id']}_{student['id']}_{q_num}", False):
                                            col1, col2 = st.columns(2)
                                            with col1:
                                                new_score = st.number_input("New Score", min_value=0, max_value=max_score, value=score, key=f"new_score_{test['id']}_{student['id']}_{q_num}")
                                            with col2:
                                                new_status = st.selectbox("New Status", ['Excellent', 'Good', 'Fair', 'Poor', 'Not Attempted'], index=['Excellent', 'Good', 'Fair', 'Poor', 'Not Attempted'].index(status) if status in ['Excellent', 'Good', 'Fair', 'Poor', 'Not Attempted'] else 0, key=f"new_status_{test['id']}_{student['id']}_{q_num}")
                                            
                                            col1, col2 = st.columns(2)
                                            if col1.button("✅ Save Score", key=f"save_score_{test['id']}_{student['id']}_{q_num}"):
                                                st.success("Score and status updated!")
                                                st.session_state[f"editing_score_{test['id']}_{student['id']}_{q_num}"] = False
                                                st.rerun()
                                            if col2.button("❌ Cancel Score", key=f"cancel_score_{test['id']}_{student['id']}_{q_num}"):
                                                st.session_state[f"editing_score_{test['id']}_{student['id']}_{q_num}"] = False
                                                st.rerun()
                                        
                                        # Show feedback with edit capability
                                        st.markdown("**Feedback:**")
                                        col1, col2 = st.columns([4, 1])
                                        with col1:
                                            st.write(qa.get('feedback', 'No feedback available.'))
                                        with col2:
                                            if st.button("✏️ Edit Feedback", key=f"edit_feedback_{test['id']}_{student['id']}_{q_num}"):
                                                st.session_state[f"editing_feedback_{test['id']}_{student['id']}_{q_num}"] = True
                                        
                                        # Edit feedback
                                        if st.session_state.get(f"editing_feedback_{test['id']}_{student['id']}_{q_num}", False):
                                            edited_feedback = st.text_area("Edit Feedback", qa.get('feedback', 'No feedback available.'), height=100, key=f"edit_feedback_text_{test['id']}_{student['id']}_{q_num}")
                                            col1, col2 = st.columns(2)
                                            if col1.button("✅ Save Feedback", key=f"save_feedback_{test['id']}_{student['id']}_{q_num}"):
                                                st.success("Feedback updated!")
                                                st.session_state[f"editing_feedback_{test['id']}_{student['id']}_{q_num}"] = False
                                                st.rerun()
                                            if col2.button("❌ Cancel Feedback", key=f"cancel_feedback_{test['id']}_{student['id']}_{q_num}"):
                                                st.session_state[f"editing_feedback_{test['id']}_{student['id']}_{q_num}"] = False
                                                st.rerun()
                                        
                                        # Edit Results and Regenerate Analysis for this question
                                        col1, col2 = st.columns(2)
                                        with col1:
                                            if st.button("🔄 Regenerate Analysis", key=f"regenerate_analysis_{test['id']}_{student['id']}_{q_num}"):
                                                st.session_state[f"regenerating_analysis_{test['id']}_{student['id']}_{q_num}"] = True
                                        with col2:
                                            if st.button("✏️ Edit Results", key=f"edit_results_{test['id']}_{student['id']}_{q_num}"):
                                                st.session_state[f"editing_results_{test['id']}_{student['id']}_{q_num}"] = True
                                        
                                        # Edit Results interface
                                        if st.session_state.get(f"editing_results_{test['id']}_{student['id']}_{q_num}", False):
                                            st.markdown("**✏️ Edit Results for Question {q_num}**")
                                            
                                            # Custom results prompt for this question
                                            custom_results_prompt = st.text_area(
                                                f"Custom Results Prompt for Q{q_num}",
                                                value=f"""Analyze the student's work for Question {q_num} and provide a comprehensive evaluation.

Question: {q_text}
Student's Answer: {relevant_text.strip() if relevant_text.strip() else "No answer provided"}

Please provide:
1. A detailed score (0 to {q_max_score} points)
2. Performance status (Excellent/Good/Fair/Poor/Not Attempted)
3. Comprehensive feedback covering strengths, errors, and improvement suggestions
4. Assessment of mathematical quality and reasoning
5. Completion status evaluation

Focus on mathematical accuracy, conceptual understanding, and provide actionable feedback.

Return your analysis in JSON format with the following structure:
{{
    "score": <score>,
    "max_score": <max_score>,
    "status": "<status>",
    "feedback": "<detailed feedback>",
    "extracted_work": "<work summary>",
    "mathematical_quality": "<quality assessment>",
    "completion_status": "<completion status>"
}}""",
                                                height=200,
                                                key=f"results_prompt_{test['id']}_{student['id']}_{q_num}"
                                            )
                                            
                                            # Show current results for reference
                                            st.markdown("**Current Results:**")
                                            st.info(f"**Score:** {qa.get('score', 0)}/{qa.get('max_score', 0)}")
                                            st.info(f"**Status:** {qa.get('status', 'Unknown')}")
                                            st.info(f"**Feedback:** {qa.get('feedback', 'No feedback available.')}")
                                            
                                            col1, col2, col3 = st.columns(3)
                                            if col1.button("🔄 Generate New Results", key=f"execute_results_{test['id']}_{student['id']}_{q_num}"):
                                                with st.spinner("🔄 Generating new results with custom prompt..."):
                                                    # Get the custom prompt
                                                    custom_prompt = st.session_state.get(f"results_prompt_{test['id']}_{student['id']}_{q_num}", "")
                                                    
                                                    # Get current student answer (respect overrides)
                                                    student_answer = get_current_answer_for_question(submission, q_num)
                                                    
                                                    # Call the regeneration function
                                                    new_analysis, message = regenerate_analysis_for_question(
                                                        q_num, custom_prompt, q_text, student_answer, test['rubric'], test['id'], student['id'], extracted_answers
                                                    )
                                                    
                                                    if new_analysis:
                                                        # Update the question analysis in the session state
                                                        submission_key = f"test_{test['id']}_{student['id']}"
                                                        if submission_key in st.session_state.submissions:
                                                            # Find and update the specific question analysis
                                                            for i, qa_item in enumerate(st.session_state.submissions[submission_key].get('question_analysis', [])):
                                                                if qa_item.get('question_number') == q_num:
                                                                    # Update with new analysis data
                                                                    st.session_state.submissions[submission_key]['question_analysis'][i].update(new_analysis)
                                                                    # Also update the submission object for immediate display
                                                                    submission['question_analysis'][i].update(new_analysis)
                                                                    break
                                                        
                                                        # Recalculate total score
                                                        total_raw_score = sum(q.get('score', 0) for q in st.session_state.submissions[submission_key].get('question_analysis', []))
                                                        total_possible_score = sum(q.get('max_score', 0) for q in st.session_state.submissions[submission_key].get('question_analysis', []))
                                                        final_percentage = math.ceil((total_raw_score / total_possible_score) * 100) if total_possible_score > 0 else 0
                                                        
                                                        # Update total score in submission
                                                        st.session_state.submissions[submission_key]['total_score'] = final_percentage
                                                        submission['total_score'] = final_percentage
                                                        
                                                        # Save to database for persistence
                                                        save_updated_results_to_database(test['id'], student['id'], q_num, new_analysis)
                                                        
                                                        st.success(f"New results generated successfully! Score updated to {final_percentage}/100")
                                                    else:
                                                        st.error(f"Failed to generate new results: {message}")
                                                    
                                                    st.session_state[f"editing_results_{test['id']}_{student['id']}_{q_num}"] = False
                                                    st.rerun()
                                            if col2.button("✅ Use Default Prompt", key=f"default_results_{test['id']}_{student['id']}_{q_num}"):
                                                with st.spinner("🔄 Generating new results with default prompt..."):
                                                    # Get current student answer (respect overrides)
                                                    student_answer = get_current_answer_for_question(submission, q_num)
                                                    
                                                    # Call the regeneration function with default prompt
                                                    new_analysis, message = regenerate_analysis_for_question(
                                                        q_num, "", q_text, student_answer, test['rubric'], test['id'], student['id'], extracted_answers
                                                    )
                                                    
                                                    if new_analysis:
                                                        # Update the question analysis in the session state
                                                        submission_key = f"test_{test['id']}_{student['id']}"
                                                        if submission_key in st.session_state.submissions:
                                                            # Find and update the specific question analysis
                                                            for i, qa_item in enumerate(st.session_state.submissions[submission_key].get('question_analysis', [])):
                                                                if qa_item.get('question_number') == q_num:
                                                                    # Update with new analysis data
                                                                    st.session_state.submissions[submission_key]['question_analysis'][i].update(new_analysis)
                                                                    # Also update the submission object for immediate display
                                                                    submission['question_analysis'][i].update(new_analysis)
                                                                    break
                                                        
                                                        # Recalculate total score
                                                        total_raw_score = sum(q.get('score', 0) for q in st.session_state.submissions[submission_key].get('question_analysis', []))
                                                        total_possible_score = sum(q.get('max_score', 0) for q in st.session_state.submissions[submission_key].get('question_analysis', []))
                                                        final_percentage = math.ceil((total_raw_score / total_possible_score) * 100) if total_possible_score > 0 else 0
                                                        
                                                        # Update total score in submission
                                                        st.session_state.submissions[submission_key]['total_score'] = final_percentage
                                                        submission['total_score'] = final_percentage
                                                        
                                                        # Save to database for persistence
                                                        save_updated_results_to_database(test['id'], student['id'], q_num, new_analysis)
                                                        
                                                        st.success(f"New results generated with default prompt! Score updated to {final_percentage}/100")
                                                    else:
                                                        st.error(f"Failed to generate new results: {message}")
                                                    
                                                    st.session_state[f"editing_results_{test['id']}_{student['id']}_{q_num}"] = False
                                                    st.rerun()
                                            if col3.button("❌ Cancel Edit", key=f"cancel_results_{test['id']}_{student['id']}_{q_num}"):
                                                st.session_state[f"editing_results_{test['id']}_{student['id']}_{q_num}"] = False
                                                st.rerun()
                                        
                                        # Analysis regeneration interface
                                        if st.session_state.get(f"regenerating_analysis_{test['id']}_{student['id']}_{q_num}", False):
                                            st.markdown("**🔄 Regenerate Analysis for Question {q_num}**")
                                            
                                            # Custom analysis prompt for this question
                                            custom_analysis_prompt = st.text_area(
                                                f"Custom Analysis Prompt for Q{q_num}",
                                                value=f"Analyze the student's work for Question {q_num}. Provide detailed feedback on mathematical accuracy, completeness, and areas for improvement.",
                                                height=100,
                                                key=f"analysis_prompt_{test['id']}_{student['id']}_{q_num}"
                                            )
                                            
                                            # Show current analysis
                                            st.markdown("**Current Analysis:**")
                                            st.info(f"**Score:** {qa.get('score', 0)}/{qa.get('max_score', 0)}")
                                            st.info(f"**Status:** {qa.get('status', 'Unknown')}")
                                            st.info(f"**Feedback:** {qa.get('feedback', 'No feedback available.')}")
                                            if qa.get('extracted_work'):
                                                st.info(f"**Work Summary:** {qa.get('extracted_work')}")
                                            
                                            col1, col2, col3 = st.columns(3)
                                            if col1.button("🔄 Regenerate with Custom Prompt", key=f"execute_analysis_{test['id']}_{student['id']}_{q_num}"):
                                                with st.spinner("🔄 Regenerating analysis with custom prompt..."):
                                                    # Get the custom prompt
                                                    custom_prompt = st.session_state.get(f"analysis_prompt_{test['id']}_{student['id']}_{q_num}", "")
                                                    
                                                    # Debug information
                                                    if st.session_state.get('debug_mode', False):
                                                        st.write(f"Debug: Custom analysis prompt: {custom_prompt}")
                                                        st.write(f"Debug: Question number: {q_num}")
                                                        st.write(f"Debug: Question text: {q_text[:100]}...")
                                                    
                                                    # Get current student answer
                                                    student_answer = extract_question_answer(extracted_answers, q_num)
                                                    
                                                    if st.session_state.get('debug_mode', False):
                                                        st.write(f"Debug: Student answer: {student_answer[:100]}...")
                                                    
                                                    # Call the regeneration function
                                                    new_analysis, message = regenerate_analysis_for_question(
                                                        q_num, custom_prompt, q_text, student_answer, test['rubric'], test['id'], student['id'], extracted_answers
                                                    )
                                                    
                                                    if new_analysis:
                                                        if st.session_state.get('debug_mode', False):
                                                            st.write(f"Debug: New analysis received: {new_analysis}")
                                                        
                                                        # Update the question analysis in the session state
                                                        submission_key = f"test_{test['id']}_{student['id']}"
                                                        if submission_key in st.session_state.submissions:
                                                            # Find and update the specific question analysis
                                                            for i, qa_item in enumerate(st.session_state.submissions[submission_key].get('question_analysis', [])):
                                                                if qa_item.get('question_number') == q_num:
                                                                    st.session_state.submissions[submission_key]['question_analysis'][i].update(new_analysis)
                                                                    if st.session_state.get('debug_mode', False):
                                                                        st.write(f"Debug: Updated question {q_num} in session state")
                                                                    break
                                                        
                                                        st.success(f"Analysis regenerated successfully! {message}")
                                                    else:
                                                        st.error(f"Failed to regenerate analysis: {message}")
                                                    
                                                    st.session_state[f"regenerating_analysis_{test['id']}_{student['id']}_{q_num}"] = False
                                                    st.rerun()
                                            if col2.button("✅ Use Default Prompt", key=f"default_analysis_{test['id']}_{student['id']}_{q_num}"):
                                                with st.spinner("🔄 Regenerating analysis with default prompt..."):
                                                    # Get current student answer
                                                    student_answer = extract_question_answer(extracted_answers, q_num)
                                                    
                                                    # Call the regeneration function with default prompt
                                                    new_analysis, message = regenerate_analysis_for_question(
                                                        q_num, "", q_text, student_answer, test['rubric'], test['id'], student['id'], extracted_answers
                                                    )
                                                    
                                                    if new_analysis:
                                                        # Update the question analysis in the session state
                                                        submission_key = f"test_{test['id']}_{student['id']}"
                                                        if submission_key in st.session_state.submissions:
                                                            # Find and update the specific question analysis
                                                            for i, qa_item in enumerate(st.session_state.submissions[submission_key].get('question_analysis', [])):
                                                                if qa_item.get('question_number') == q_num:
                                                                    st.session_state.submissions[submission_key]['question_analysis'][i].update(new_analysis)
                                                                    break
                                                        
                                                        st.success(f"Analysis regenerated with default prompt! {message}")
                                                    else:
                                                        st.error(f"Failed to regenerate analysis: {message}")
                                                    
                                                    st.session_state[f"regenerating_analysis_{test['id']}_{student['id']}_{q_num}"] = False
                                                    st.rerun()
                                            if col3.button("❌ Cancel Analysis", key=f"cancel_analysis_{test['id']}_{student['id']}_{q_num}"):
                                                st.session_state[f"regenerating_analysis_{test['id']}_{student['id']}_{q_num}"] = False
                                                st.rerun()
                                        
                                        # Show mathematical quality assessment if available
                                        if qa.get('mathematical_quality'):
                                            st.info(f"**Mathematical Quality:** {qa.get('mathematical_quality')}")
                                        
                                        # Show completion status if available
                                        if qa.get('completion_status'):
                                            completion_status = qa.get('completion_status')
                                            completion_color = {
                                                'Complete': 'success',
                                                'Nearly Complete': 'info',
                                                'Partially Complete': 'warning',
                                                'Incomplete': 'error',
                                                'Abandoned': 'error'
                                            }.get(completion_status, 'secondary')
                                            st.write(f"**Completion:** :{completion_color}[{completion_status}]")
                                    else:
                                        st.warning("No analysis available for this question.")
                                    
                                    st.divider()
                        else:
                            st.warning("Could not parse questions from the test.")

                        st.markdown("**📊 Detailed Question Analysis**")
                        question_analysis = submission.get('question_analysis', [])
                        if question_analysis:
                            for qa in question_analysis:
                                with st.container(border=True):
                                    q_col1, q_col2 = st.columns([3,1])
                                    q_col1.write(f"**Question {qa.get('question_number', 'N/A')}**")
                                    q_col2.metric("Score", f"{qa.get('score', 0)}/{qa.get('max_score', 0)}")
                                    
                                    # Show extracted work if available
                                    if qa.get('extracted_work'):
                                        st.info(f"**Work Found:** {qa.get('extracted_work')}")
                                    
                                    # Show status and feedback
                                    status = qa.get('status', 'Unknown')
                                    status_color = {
                                        'Excellent': 'success',
                                        'Good': 'info', 
                                        'Fair': 'warning',
                                        'Poor': 'error',
                                        'Not Attempted': 'error'
                                    }.get(status, 'secondary')
                                    
                                    st.write(f"**Status:** :{status_color}[{status}]")
                                    st.write(f"**Feedback:** {qa.get('feedback', 'No feedback available.')}")
                                    
                                    # Show mathematical quality assessment if available
                                    if qa.get('mathematical_quality'):
                                        st.info(f"**Mathematical Quality:** {qa.get('mathematical_quality')}")
                                    
                                    # Show completion status if available
                                    if qa.get('completion_status'):
                                        completion_status = qa.get('completion_status')
                                        completion_color = {
                                            'Complete': 'success',
                                            'Nearly Complete': 'info',
                                            'Partially Complete': 'warning',
                                            'Incomplete': 'error',
                                            'Abandoned': 'error'
                                        }.get(completion_status, 'secondary')
                                        st.write(f"**Completion:** :{completion_color}[{completion_status}]")
                        else:
                            st.caption("No question-specific analysis was generated.")
                        # Completed Rubric and Summary Tables
                        st.markdown("**Rubric Evaluation**")
                        cached_rows = st.session_state.get(f"rubric_rows_{test['id']}")
                        if cached_rows:
                            completed = build_completed_rubric_table(cached_rows, questions, submission.get('question_analysis', []))
                            if completed:
                                st.markdown("Completed Rubric Table")
                                st.table(completed)
                                st.markdown("Summary for Teacher")
                                summary_rows = generate_rubric_summary_table(completed)
                                if summary_rows:
                                    st.table(summary_rows)
                                else:
                                    st.caption("No summary available.")
                            else:
                                st.caption("Rubric present, but could not compute evaluation.")
                        else:
                            st.caption("No rubric uploaded for this test.")
                        
                        # Show additional analysis if available
                        if submission.get('mathematical_thinking'):
                            st.markdown("**Mathematical Thinking Assessment**")
                            st.info(submission.get('mathematical_thinking'))
                        
                        if submission.get('learning_recommendations'):
                            st.markdown("**Learning Recommendations**")
                            st.success(submission.get('learning_recommendations'))
                        
                        # Show saved file information
                        if submission.get('saved_file'):
                            st.divider()
                            st.markdown("**📁 Extracted Content Saved**")
                            st.info(f"All extracted content has been saved to: `{submission.get('saved_file')}`")
                            
                            # Provide download button for the saved file
                            try:
                                with open(submission.get('saved_file'), 'rb') as f:
                                    file_content = f.read()
                                
                                st.download_button(
                                    label="📥 Download Complete Report (DOC)",
                                    data=file_content,
                                    file_name=f"grading_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    help="Download the complete grading report as a DOC file with all extracted content",
                                    key=f"download_doc_{test['id']}_{student['id']}"
                                )
                            except Exception as e:
                                st.warning(f"Could not prepare download: {e}")
                else:
                    cols[1].info("Pending")
                    uploaded_files = cols[2].file_uploader(
                        "Upload Handwritten Answers", 
                        type=["pdf", "png", "jpg", "jpeg"], 
                        accept_multiple_files=True,
                        key=f"uploader_{submission_key}",
                        help="Upload PDF files or image files (PNG, JPG, JPEG). You can upload multiple files."
                    )
                    if uploaded_files:
                        progress_bar = cols[2].progress(0, text="Processing files...")
                        answer_images = process_uploaded_files(uploaded_files)
                        
                        # Display processing information
                        st.subheader("📸 Answer Script Processing:")
                        st.info(f"📁 Processing {len(uploaded_files)} uploaded file(s)")
                        
                        if answer_images:
                            st.success(f"✅ Successfully converted {len(answer_images)} image(s) for analysis")
                            
                            # Store original images in session state for regeneration
                            st.session_state[f"original_images_{test['id']}_{student['id']}"] = answer_images
                            
                            # Show preview of processed images
                            with st.expander("👁️ Preview Processed Answer Images", expanded=False):
                                for i, img in enumerate(answer_images):
                                    st.write(f"**Image {i+1}:**")
                                    st.image(img, caption=f"Answer Image {i+1}", use_container_width=True)
                                    
                                    # Extract and display text from each image
                                    extracted_text = extract_text_from_image(img)
                                    if extracted_text and not extracted_text.startswith("GPT-4o OCR failed"):
                                        st.text_area(f"Extracted Text from Image {i+1}", extracted_text, height=100, disabled=True)
                                        st.info(f"✅ GPT-4o extracted {len(extracted_text)} characters from Image {i+1}")
                                    else:
                                        st.warning(f"⚠️ {extracted_text}")
                                        st.info("💡 **Note:** The AI grading will still work using image analysis, even without text extraction.")
                                    st.divider()
                            
                            progress_bar.progress(0.1, text="Starting evaluation...")
                            result = grade_handwritten_submission_with_gpt4o(test['question_text'], answer_images, test['rubric'], test['title'], progress_bar)
                            
                            # Debug: Check if extracted answers are in result
                            if st.session_state.get('debug_mode', False):
                                st.subheader("🔍 Debug: Result Contents")
                                st.write("Result keys:", list(result.keys()) if isinstance(result, dict) else "Not a dict")
                                if isinstance(result, dict) and 'extracted_answers' in result:
                                    st.write("Extracted answers found:", len(result['extracted_answers']))
                                else:
                                    st.write("No extracted_answers in result")
                            
                            st.session_state.submissions[submission_key] = {"status": "error" if "error" in result else "graded", **result}
                            st.rerun()
                        else:
                            st.error("No valid files could be processed. Please check your uploads.")
